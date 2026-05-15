import io
import os
import docx
from docx import Document


class DocParser:
    def __init__(self):
        pass

    def parse(self, file_path_or_bytes, return_position_map=False):
        """
        解析 docx 文件，返回清洗后的单行长文本
        
        Args:
            file_path_or_bytes: 文件路径或二进制流
            return_position_map: 是否返回位置映射（字符位置 -> (row, column)）
        
        Returns:
            如果 return_position_map=False: 返回文本字符串
            如果 return_position_map=True: 返回 (text, position_map) 元组
                position_map: dict，key 是字符在文本中的位置，value 是 (row, column) 元组
        """
        # 兼容文件路径或二进制流
        if isinstance(file_path_or_bytes, str):
            # 检查文件是否存在
            if not os.path.exists(file_path_or_bytes):
                raise FileNotFoundError(f"文件不存在: {file_path_or_bytes}")

            # 尝试作为 docx 文件打开，如果失败则作为纯文本处理
            try:
                doc = Document(file_path_or_bytes)
            except Exception as e:
                # 如果不是有效的 docx 文件，尝试作为纯文本读取
                if "Package not found" in str(e) or "not a zipfile" in str(e).lower():
                    with open(file_path_or_bytes, 'r', encoding='utf-8') as f:
                        text = f.read()
                    # 按行分割并清理，去除每行末尾的句号避免重复
                    lines = []
                    for line in text.split('\n'):
                        line = line.strip()
                        if line:
                            # 去除末尾的句号（如果存在），避免连接时重复
                            if line.endswith('。'):
                                line = line[:-1]
                            lines.append(line)
                    # 使用换行符连接，保留段落结构
                    result = "\n".join(lines)
                    
                    if return_position_map:
                        # 构建位置映射
                        position_map = {}
                        current_pos = 0
                        for row, line in enumerate(lines):
                            line_len = len(line)
                            for col in range(line_len):
                                position_map[current_pos] = (row, col)
                                current_pos += 1
                            # 换行符位置
                            if row < len(lines) - 1:  # 不是最后一行
                                position_map[current_pos] = (row, line_len)
                                current_pos += 1
                        return result, position_map
                    
                    return result
                else:
                    raise
        else:
            doc = Document(io.BytesIO(file_path_or_bytes))

        full_text_list = []

        # 1. 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 去除末尾的句号（如果存在），避免连接时重复
                if text.endswith('。'):
                    text = text[:-1]
                full_text_list.append(text)

        # 2. 提取表格 (线性化处理：保留语义的关键)
        for table in doc.tables:
            for row in table.rows:
                # 将表格的一行转换为 "Key: Value" 或 "A B C" 的形式
                cells = [cell.text.strip()
                         for cell in row.cells if cell.text.strip()]
                if cells:
                    # 用空格连接单元格，形成自然语句
                    row_text = " ".join(cells)
                    # 去除末尾的句号（如果存在）
                    if row_text.endswith('。'):
                        row_text = row_text[:-1]
                    full_text_list.append(row_text)

        # 拼接为一个长字符串，使用换行符连接以保留段落结构
        result = "\n".join(full_text_list)
        
        if return_position_map:
            # 构建位置映射：字符位置 -> (row, column)
            # 注意：这里假设文本是用 \n 连接的，每个段落是一行
            position_map = {}
            current_pos = 0
            
            for row, line in enumerate(full_text_list):
                line_len = len(line)
                # 记录该行的每个字符位置
                for col in range(line_len):
                    position_map[current_pos] = (row, col)
                    current_pos += 1
                # 换行符位置
                if row < len(full_text_list) - 1:  # 不是最后一行
                    position_map[current_pos] = (row, line_len)  # 换行符位置
                    current_pos += 1
            
            return result, position_map
        
        return result


# --- 单元测试 ---
if __name__ == "__main__":
    # 您需要创建一个名为 test.docx 的文件来测试
    parser = DocParser()
    # 使用相对于脚本文件位置的路径，这样无论从哪里运行都能正确找到文件
    script_dir = os.path.dirname(os.path.abspath(__file__))
    test_file = os.path.join(script_dir, "test.docx")
    print(parser.parse(test_file))
    # pass
