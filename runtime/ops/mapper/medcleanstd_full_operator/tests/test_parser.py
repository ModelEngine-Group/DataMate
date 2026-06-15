from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from myparser.parser import DocParser


class TestDocParser:
    def test_parse_text_file_keeps_non_empty_lines(self, tmp_path: Path):
        text_path = tmp_path / "sample.txt"
        text_path.write_text("line1\n\nline2\n", encoding="utf-8")

        parser = DocParser()
        parsed = parser.parse(str(text_path))

        assert parsed == "line1\nline2"

    def test_parse_docx_file_reads_paragraphs_and_table(self, tmp_path: Path):
        docx_path = tmp_path / "sample.docx"
        doc = Document()
        doc.add_paragraph("paragraph one")
        doc.add_paragraph("paragraph two")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "key"
        table.cell(0, 1).text = "value"
        doc.save(str(docx_path))

        parser = DocParser()
        parsed = parser.parse(str(docx_path))

        assert "paragraph one" in parsed
        assert "paragraph two" in parsed
        assert "key value" in parsed

    def test_parse_docx_bytes(self):
        buffer = io.BytesIO()
        doc = Document()
        doc.add_paragraph("from-bytes")
        doc.save(buffer)

        parser = DocParser()
        parsed = parser.parse(buffer.getvalue())

        assert parsed == "from-bytes"

    def test_parse_with_position_map(self, tmp_path: Path):
        text_path = tmp_path / "map.txt"
        text_path.write_text("ab\ncd", encoding="utf-8")

        parser = DocParser()
        parsed, position_map = parser.parse(str(text_path), return_position_map=True)

        assert parsed == "ab\ncd"
        assert position_map[0] == (0, 0)
        assert position_map[2] == (0, 2)
        assert position_map[3] == (1, 0)
