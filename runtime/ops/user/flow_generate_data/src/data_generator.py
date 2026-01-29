"""
数据生成模块 - 用于生成流水资产分析表的随机数据
"""

import os
import re
import json
from docx import Document
from docx.shared import Pt
from faker import Faker
from datetime import datetime, timedelta
from typing import Dict, Any, List, Optional


class DataGenerator:
    """数据生成器类"""

    def __init__(self, seed: Optional[int] = None):
        """
        初始化数据生成器

        Args:
            seed: 随机种子，用于控制生成的一致性
        """
        self.fake = Faker('zh_CN')
        if seed is not None:
            Faker.seed(seed)

    def generate_values(self) -> Dict[str, Any]:
        """
        生成所有需要的填充值

        Returns:
            包含所有字段值的字典
        """
        # 生成姓名
        户名 = self.fake.name()
        经办人 = self.fake.name()
        复核人 = self.fake.name()

        # 生成银行卡号
        银行列表 = ["洛汀城市银行", "工商银行", "农业银行", "中国银行", "建设银行", 
                   "交通银行", "招商银行", "浦发银行", "中信银行", "光大银行", 
                   "民生银行", "平安银行", "华夏银行", "兴业银行", "广发银行"]
        银行名称 = self.fake.random_element(银行列表)
        账号 = 银行名称 + self.fake.credit_card_number(card_type=None)[:16]

        # 生成统计期间（6个月）
        start_date = self.fake.date_between(start_date='-1y', end_date='today')
        end_date = start_date + timedelta(days=180)
        统计期间 = f"{start_date.strftime('%Y年%m月%d日')}-{end_date.strftime('%Y年%m月%d日')}"

        # 生成日期
        today = datetime.now()
        日期年份 = str(today.year)
        日期月份 = f"{today.month:02d}"
        日期日 = f"{today.day:02d}"

        # 生成数字和金额
        流入资金笔数 = str(self.fake.random_int(5, 50))
        流入资金总额 = str(self.fake.random_int(5000, 50000))
        月平均流入资金 = str(int(int(流入资金总额) / 6))

        # 每月还款金额
        每月还款金额 = f"{self.fake.random.uniform(1000, 5000):.2f}"

        # 目前月债务支出
        目前月债务支出 = str(self.fake.random_int(0, 3000))

        # 借款人家庭当前月总收入
        借款人家庭当前月总收入 = f"{self.fake.random.uniform(5000, 15000):.2f}"

        # 计算债务占比
        total_debt = float(每月还款金额) + float(目前月债务支出)
        债务占比 = f"{(total_debt / float(借款人家庭当前月总收入) * 100):.2f}%"

        # 生成备注
        备注选项 = ['工资收入', '租金流入', '经营收入', '投资收益', '其他收入']
        备注 = self.fake.random_element(elements=备注选项)

        # 返回所有填充值
        return {
            '户名': 户名,
            '账号': 账号,
            '统计期间': 统计期间,
            '流入资金笔数': 流入资金笔数,
            '流入资金总额': 流入资金总额,
            '月平均流入资金': 月平均流入资金,
            '备注': 备注,
            '每月还款金额': 每月还款金额,
            '目前月债务支出': 目前月债务支出,
            '借款人家庭当前月总收入': 借款人家庭当前月总收入,
            '债务占比': 债务占比,
            '经办人': 经办人,
            '复核人': 复核人,
            '日期年份': 日期年份,
            '日期月份': 日期月份,
            '日期日': 日期日,
            '每月还款金额下划线': '',
            '目前月债务支出下划线': ''
        }

    def _adjust_underline_spaces(self, values: Dict[str, Any]) -> Dict[str, Any]:
        """
        调整下划线空格数量，使排版更美观

        Args:
            values: 原始值字典

        Returns:
            调整后的值字典
        """
        # 调整经办人和复核人的下划线长度，使其一致
        经办人长度 = len(values['经办人'])
        复核人长度 = len(values['复核人'])

        # 基础空格数量为6个（相当于6个下划线）
        基础空格 = " " * 6

        # 计算差值
        差值 = abs(经办人长度 - 复核人长度)

        # 根据差值调整空格数量
        if 经办人长度 < 复核人长度:
            # 经办人字数少，需要增加空格
            经办人下划线 = 基础空格 + " " * 差值 + values['经办人'] + 基础空格 + " " * 差值
            复核人下划线 = 基础空格 + values['复核人'] + 基础空格
        elif 复核人长度 < 经办人长度:
            # 复核人字数少，需要增加空格
            经办人下划线 = 基础空格 + values['经办人'] + 基础空格
            复核人下划线 = 基础空格 + " " * 差值 + values['复核人'] + 基础空格 + " " * 差值
        else:
            # 字数相同，使用相同的空格数量
            经办人下划线 = 基础空格 + values['经办人'] + 基础空格
            复核人下划线 = 基础空格 + values['复核人'] + 基础空格

        values['经办人下划线'] = 经办人下划线
        values['复核人下划线'] = 复核人下划线

        # 调整每月还款金额和目前月债务支出的下划线
        每月还款金额_str = str(values['每月还款金额'])
        目前月债务支出_str = str(values['目前月债务支出'])

        # 计算每月还款金额的下划线数量
        每月还款金额长度 = len(每月还款金额_str)
        每月还款金额前空格 = "   "  # 3个空格（相当于3个下划线）
        每月还款金额后空格 = "   "  # 3个空格（相当于3个下划线）

        # 计算目前月债务支出的下划线数量
        目前月债务支出长度 = len(目前月债务支出_str)
        目前月债务支出前空格 = "   "  # 3个空格（相当于3个下划线）
        目前月债务支出后空格 = "   "  # 3个空格（相当于3个下划线）

        # 统一前后3个空格，不进行差额计算

        values['每月还款金额下划线'] = 每月还款金额前空格 + 每月还款金额_str + 每月还款金额后空格
        values['目前月债务支出下划线'] = 目前月债务支出前空格 + 目前月债务支出_str + 目前月债务支出后空格
        values['借款人家庭当前月总收入下划线'] = '   ' + str(values['借款人家庭当前月总收入']) + '   '
        values['债务占比下划线'] = '   ' + str(values['债务占比']) + '   '

        return values

    def _extract_fields_from_text(self, text: str) -> List[str]:
        """
        从文本中提取 {{字段名}} 格式的字段

        Args:
            text: 文本内容

        Returns:
            字段名列表
        """
        pattern = r'\{\{(.*?)\}\}'
        matches = re.findall(pattern, text)
        return [match.strip() for match in matches]

    def _extract_placeholder_fields(self, template_path: str) -> List[str]:
        """
        从 Word 模板中提取字段

        Args:
            template_path: Word模板文件路径

        Returns:
            字段名列表
        """
        doc = Document(template_path)
        fields_set = set()

        # 在段落中查找
        for para in doc.paragraphs:
            fields_set.update(self._extract_fields_from_text(para.text))

        # 在表格中查找
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        fields_set.update(self._extract_fields_from_text(para.text))

        return list(fields_set)

    def fill_document(self, template_path: str, output_path: str, 
                     font_name: str = '宋体', font_size: int = 12) -> Dict[str, Any]:
        """
        填充Word文档

        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            font_name: 字体名称
            font_size: 字体大小

        Returns:
            填充的值字典
        """
        # 生成填充值
        values = self.generate_values()

        # 调整下划线空格
        values = self._adjust_underline_spaces(values)

        # 加载模板文档
        doc = Document(template_path)

        # 处理所有段落
        for paragraph in doc.paragraphs:
            self._fill_paragraph(paragraph, values, font_name, font_size)

        # 处理所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._fill_paragraph(paragraph, values, font_name, font_size)

        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # 保存填充后的文档
        doc.save(output_path)

        # 保存填充值为JSON文件
        json_output_file = output_path.replace('.docx', '.json')
        with open(json_output_file, 'w', encoding='utf-8') as f:
            json.dump(values, f, ensure_ascii=False, indent=2)

        return values

    def _fill_paragraph(self, paragraph, values: Dict[str, Any], 
                      font_name: str, font_size: int):
        """
        填充段落中的占位符

        Args:
            paragraph: Word段落对象
            values: 值字典
            font_name: 字体名称
            font_size: 字体大小
        """
        # 查找所有占位符
        pattern = r'\{\{(.*?)\}\}'
        matches = list(re.finditer(pattern, paragraph.text))

        if matches:
            # 保存原始文本
            original_text = paragraph.text

            # 清空段落
            paragraph.clear()

            # 重建段落内容
            last_end = 0
            for match in matches:
                # 添加占位符前的文本
                if match.start() > last_end:
                    pre_text = original_text[last_end:match.start()]
                    if pre_text:
                        run = paragraph.add_run(pre_text)
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        run.font.size = Pt(font_size)

                # 处理占位符
                field_name = match.group(1)
                field_value = self._get_field_value(field_name, values, paragraph.text)

                # 如果没有找到对应的值，则使用原始字段名
                if field_value is None:
                    field_value = field_name

                # 添加填充的值
                underline = self._check_underline(field_name)
                run = paragraph.add_run(field_value)
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = Pt(font_size)
                if underline:
                    run.font.underline = True

                last_end = match.end()

            # 添加最后的文本
            if last_end < len(original_text):
                post_text = original_text[last_end:]
                if post_text:
                    run = paragraph.add_run(post_text)
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(font_size)

    def _get_field_value(self, field_name: str, values: Dict[str, Any], 
                        context_text: str) -> Optional[str]:
        """
        根据字段名获取对应的值

        Args:
            field_name: 字段名
            values: 值字典
            context_text: 上下文文本

        Returns:
            字段值，如果未找到则返回None
        """
        # 根据字段名中的特征匹配对应的值
        if '张三' in field_name:
            return values['户名']
        elif '洛汀城市银行' in field_name:
            return values['账号']
        elif '2025年01月01日-2025年06月30日' in field_name:
            return values['统计期间']
        elif '27' in field_name and '笔数' in context_text:
            return values['流入资金笔数']
        elif '2781' in field_name and '总额' in context_text:
            return values['流入资金总额']
        elif '4651' in field_name and '平均' in context_text:
            return values['月平均流入资金']
        elif '租金流入' in field_name:
            return values['备注']
        elif '1860.08' in field_name:
            return values['每月还款金额下划线']
        elif '0' in field_name and '债务' in context_text:
            return values['目前月债务支出下划线']
        elif '4789.05' in field_name:
            return values['借款人家庭当前月总收入下划线']
        elif '38.74%' in field_name:
            return values['债务占比下划线']
        elif '金某' in field_name:
            return values['经办人下划线']
        elif '黄某' in field_name:
            return values['复核人下划线']
        elif '2025' in field_name and '年' in context_text:
            return values['日期年份']
        elif '06' in field_name and '月' in context_text:
            return values['日期月份']
        elif '05' in field_name and '日' in context_text:
            return values['日期日']

        return None

    def _check_underline(self, field_name: str) -> bool:
        """
        检查字段是否需要添加下划线

        Args:
            field_name: 字段名

        Returns:
            是否需要下划线
        """
        return field_name in ['金某', '黄某', '1860.08', '0', '4789.05', '38.74%']


def qn(tag):
    """获取命名空间的限定名"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    return qn(tag)
