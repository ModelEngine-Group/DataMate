"""
文档生成模块 - 最终修复版
"""

import os
import shutil
from docx import Document
from typing import Dict, Any, List
from pathlib import Path


class DocGenerator:
    """文档生成器 - 使用{{}}占位符"""

    def __init__(self, template_path: str, output_dir: str = "generated_docs"):
        self.template_path = template_path
        self.output_dir = output_dir
        Path(output_dir).mkdir(parents=True, exist_ok=True)

    def _replace_in_runs(self, runs, data: Dict[str, Any]):
        """在runs列表中替换所有占位符

        处理 {{field_name}} 格式，字段名可能跨多个 run
        例如: {{联系人姓名}} 可能被拆分为 {{ + 联系人 + 姓名 + }}
        """
        i = 0
        while i < len(runs):
            # 查找 {{
            if '{{' in runs[i].text:
                # 向后查找 }}
                j = i + 1
                field_name_parts = []
                end_run_idx = -1

                while j < len(runs):
                    if '}}' in runs[j].text:
                        end_run_idx = j
                        break
                    # 收集中间的文本作为字段名
                    if runs[j].text.strip():
                        field_name_parts.append(runs[j].text.strip())
                    j += 1

                # 如果找到了完整的 {{}} 模式
                if end_run_idx != -1 and field_name_parts:
                    field_name = ''.join(field_name_parts)

                    # 从data中获取值
                    if field_name in data:
                        value = str(data[field_name])

                        # 替换逻辑：
                        # 1. 清理 run[i] 中的 {{
                        runs[i].text = runs[i].text.replace('{{', '')

                        # 2. 清空所有中间 run 的文本
                        for k in range(i + 1, end_run_idx):
                            runs[k].text = ''

                        # 3. 在第一个中间 run 设置实际值
                        if i + 1 < end_run_idx:
                            runs[i + 1].text = value

                        # 4. 清理 run[end_run_idx] 中的 }}
                        runs[end_run_idx].text = runs[end_run_idx].text.replace('}}', '')

                    # 移动到 }} 之后
                    i = end_run_idx + 1
                    continue

            i += 1

    def generate(self, data: Dict[str, Any], output_name: str) -> str:
        """根据数据生成文档"""
        output_path = os.path.join(self.output_dir, f"{output_name}.docx")
        shutil.copy(self.template_path, output_path)

        doc = Document(output_path)

        # 替换段落中的占位符
        for para in doc.paragraphs:
            self._replace_in_runs(para.runs, data)

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_runs(para.runs, data)

        doc.save(output_path)
        return output_path

    def generate_batch(self, records: List[Dict[str, Any]], prefix: str = "loan_clearance") -> List[str]:
        """批量生成文档"""
        output_paths = []
        for i, record in enumerate(records):
            index_str = f"{i + 1:03d}"
            output_name = f"{prefix}_{index_str}"
            try:
                path = self.generate(record, output_name)
                output_paths.append(path)
                print(f"  已生成: {output_name}.docx")
            except Exception as e:
                print(f"  生成失败 {output_name}: {e}")
        return output_paths
