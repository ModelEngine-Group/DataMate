#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
收入证明文档处理工具
功能：从模板文件中提取标记字段，生成回填数据，并填充到模板中生成新文档
"""

import re
import json
import os
from faker import Faker
from docx import Document

from loguru import logger

# 创建中文faker实例
fake = Faker('zh_CN')


# ============================================================
# 第一部分：字段提取功能
# ============================================================

def extract_field_info(docx_path):
    """
    提取docx文档中{{字段名}}标记的内容及其下划线信息
    同时处理正文段落、表格单元格中的占位符

    Args:
        docx_path: docx文件路径

    Returns:
        list: 包含字段信息的字典列表
    """
    doc = Document(docx_path)
    fields_info = []

    # 正则模式匹配{{字段名}}
    pattern = re.compile(r'\{\{([^}]+)\}\}')

    # ========== 1. 处理正文段落 ==========
    for para_idx, paragraph in enumerate(doc.paragraphs):
        _extract_from_paragraph(
            paragraph,
            fields_info,
            "body",
            {"paragraph_index": para_idx},
            pattern
        )

    # ========== 2. 处理表格（支持嵌套表格） ==========
    _process_tables_recursively(
        doc.tables,
        fields_info,
        pattern,
        table_path=[]
    )

    return fields_info


def _process_tables_recursively(tables, fields_info, pattern, table_path, parent_context=None):
    """递归处理表格及其嵌套表格"""
    for table_idx, table in enumerate(tables):
        if not table._cells:
            logger.info(f"警告: 表格{'.'.join(map(str, table_path + [table_idx]))}为空，跳过处理")
            continue

        current_table_path = table_path + [table_idx]

        # 使用 _cells 直接遍历所有单元格
        for cell_idx, cell in enumerate(table._cells):
            cell_context = {
                'table_path': current_table_path,
                'cell_index': cell_idx
            }

            if parent_context:
                cell_context.update(parent_context)

            # 处理单元格内的段落
            if cell.paragraphs:
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    _extract_from_paragraph(
                        paragraph,
                        fields_info,
                        "nested_table" if table_path else "table",
                        {
                            **cell_context,
                            'paragraph_index': para_idx
                        },
                        pattern
                    )

            # 处理单元格内的嵌套表格
            if hasattr(cell, 'tables') and cell.tables:
                _process_tables_recursively(
                    cell.tables,
                    fields_info,
                    pattern,
                    current_table_path,
                    cell_context
                )


def _extract_from_paragraph(paragraph, fields_info, location_type, context, pattern):
    """从单个段落中提取占位符（辅助函数）"""
    para_text = paragraph.text
    matches = pattern.finditer(para_text)

    for match in matches:
        field_name = match.group(1)
        start_pos = match.start()
        end_pos = match.end()

        # 检查该字段是否有下划线
        has_underline = check_underline_in_runs(paragraph, start_pos, end_pos)

        fields_info.append({
            'field_name': field_name.strip(),
            'full_match': match.group(0),
            'location_type': location_type,
            'context': context,
            'has_underline': has_underline,
            'text_context': para_text.strip()
        })


def check_underline_in_runs(paragraph, start_pos, end_pos):
    """检查段落中指定位置的文本是否有下划线"""
    current_pos = 0
    has_underline = False

    for run in paragraph.runs:
        run_start = current_pos
        run_end = current_pos + len(run.text)

        if run_end > start_pos and run_start < end_pos:
            if run.underline:
                has_underline = True
                break

        current_pos = run_end

    return has_underline


def save_fields_to_json(fields_info):
    """将提取的字段信息转换为JSON格式"""
    # 去重字段（按字段名）
    unique_fields = {}
    for field in fields_info:
        field_name = field['field_name']
        if field_name not in unique_fields:
            unique_fields[field_name] = {
                '字段名': field_name,
                '完整标记': field['full_match'],
                '位置类型': field['location_type'],
                '上下文': field['context'],
                '有下划线': field['has_underline']
            }

    # 转换为列表并排序
    fields_list = list(unique_fields.values())
    fields_list.sort(key=lambda x: x['字段名'])

    # 创建输出数据
    output_data = {
        '字段列表': fields_list
    }

    return output_data


# ============================================================
# 第二部分：数据生成功能
# ============================================================

def load_fields_info(fields_data):
    """
    从字段数据中加载字段信息

    Args:
        fields_data: 字段数据字典（可以直接来自 extract_field_info）

    Returns:
        dict: 字段信息字典
    """
    # 如果是文件路径，从文件加载
    if isinstance(fields_data, str):
        if not os.path.exists(fields_data):
            raise FileNotFoundError(f"字段文件不存在: {fields_data}")
        with open(fields_data, 'r', encoding='utf-8') as f:
            return json.load(f)

    # 如果已经是字典，直接返回
    return fields_data


def detect_field_type(field_name):
    """根据字段名推断字段类型"""
    # 姓名相关
    if field_name == '姓名':
        return 'name'

    # 身份证号
    if field_name == '身份证号':
        return 'id_card'

    # 日期相关
    if field_name == '入职日期':
        return 'date'
    if field_name == '日期':
        return 'day'

    # 年份
    if field_name == '年份':
        return 'year'

    # 月份
    if field_name == '月份':
        return 'month'

    # 职位/职务
    if field_name in ['职位', '职务']:
        return 'job'

    # 公司相关
    if field_name == '公司名称':
        return 'company_name'
    if field_name == '公司地址':
        return 'company_address'

    # 联系人
    if field_name == '联系人':
        return 'contact_name'

    # 电话
    if field_name == '联系电话':
        return 'phone'

    # 收入相关
    if '月收入' in field_name or '月均收入' in field_name:
        if '大写' in field_name:
            return 'income_uppercase'
        if '小写' in field_name:
            return 'income_lowercase'
        return 'income'
    # 年收入相关
    if '年收入' in field_name:
        if '大写' in field_name:
            return 'annual_income_uppercase'
        if '小写' in field_name:
            return 'annual_income_lowercase'
        return 'annual_income'

    # 备注
    if field_name == '备注':
        return 'remark'

    # 默认返回文本
    return 'text'


def generate_value_by_type(field_type, monthly_income=None, year_value=None, month_value=None):
    """根据字段类型生成对应的值"""
    if field_type == 'name':
        last_name = fake.last_name()
        first_name = fake.first_name()
        return f"{last_name}{first_name}"

    elif field_type == 'id_card':
        return fake.ssn()

    elif field_type == 'date':
        return fake.date_between(start_date='-5y', end_date='today').strftime('%Y-%m-%d')

    elif field_type == 'year':
        year = fake.date_between(start_date='-1y', end_date='+1y').year
        return str(year)

    elif field_type == 'month':
        return str(fake.random_int(min=1, max=12))

    elif field_type == 'day':
        import calendar
        if year_value is not None and month_value is not None:
            try:
                year = int(year_value)
                month = int(month_value)
                days_in_month = calendar.monthrange(year, month)[1]
                day = fake.random_int(min=1, max=days_in_month)
                return str(day)
            except (ValueError, IndexError):
                return str(fake.random_int(min=1, max=28))
        else:
            return str(fake.random_int(min=1, max=28))

    elif field_type == 'job':
        jobs = ['软件工程师', '产品经理', '销售经理', '人事专员',
                '市场专员', '财务经理', '行政主管', '运营经理',
                '测试工程师', 'UI设计师', '数据分析师']
        return fake.random_element(jobs)

    elif field_type == 'income_uppercase' or field_type == 'income_average_uppercase':
        income = fake.random_int(min=5000, max=50000)
        return number_to_chinese(income)

    elif field_type == 'income_lowercase' or field_type == 'income_average_lowercase':
        income = fake.random_int(min=5000, max=50000)
        return f"{income}元"

    elif field_type == 'company_name':
        return fake.company()

    elif field_type == 'company_address':
        return fake.address().replace('\n', '')

    elif field_type == 'contact_name':
        last_name = fake.last_name()
        first_name = fake.first_name()
        return f"{last_name}{first_name}"

    elif field_type == 'phone':
        if fake.boolean(chance_of_getting_true=50):
            return fake.phone_number()
        else:
            area_code = fake.random_element(['010', '021', '020', '0755', '023'])
            number = fake.numerify('#######')
            return f"{area_code}-{number}"

    elif field_type == 'annual_income_uppercase':
        if monthly_income is None:
            monthly_income = fake.random_int(min=5000, max=50000)
        bonus = fake.random_int(min=0, max=50000)
        annual_income = monthly_income * 12 + bonus
        return number_to_chinese(annual_income)

    elif field_type == 'annual_income_lowercase':
        if monthly_income is None:
            monthly_income = fake.random_int(min=5000, max=50000)
        bonus = fake.random_int(min=0, max=50000)
        annual_income = monthly_income * 12 + bonus
        return f"{annual_income}元"

    elif field_type == 'remark':
        remarks = [
            '该员工在职期间表现优秀，收入稳定可靠。',
            '特此证明该员工的收入情况，供相关单位参考。',
            '该员工为本单位正式员工，月收入情况属实。',
            '此证明仅用于证明该员工的收入情况，其他无效。',
            '该员工工作认真负责，收入水平与岗位相符。'
        ]
        return fake.random_element(remarks)

    else:
        return fake.word()


def number_to_chinese(num):
    """将数字转换为中文大写"""
    digits = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    units = ['', '拾', '佰', '仟', '万', '拾', '佰', '仟', '亿']

    if num == 0:
        return '零元整'

    result = ''
    num_str = str(num)
    length = len(num_str)

    for i, digit in enumerate(num_str):
        digit_int = int(digit)
        pos = length - i - 1

        if digit_int != 0:
            result += digits[digit_int] + units[pos]
        else:
            if result and result[-1] != '零':
                result += '零'

    if result.endswith('零'):
        result = result[:-1]

    return result + '元整'


def calculate_work_months(start_date_str, end_date_str):
    """计算两个日期之间的工作月数"""
    from datetime import datetime

    start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
    end_date = datetime.strptime(end_date_str, '%Y-%m-%d')

    months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)

    if end_date.day < start_date.day:
        months -= 1

    return max(1, months)


def generate_fill_data(template_path):
    """
    从模板文件中提取字段并生成回填数据

    Args:
        template_path: 模板文档路径

    Returns:
        dict: 包含所有字段和生成值的字典
    """
    from datetime import datetime, timedelta

    logger.info(f"正在处理模板文件: {template_path}\n")

    # ========== 第一步：从模板中提取字段 ==========
    logger.info("=== 第一步：提取字段信息 ===")
    fields_info = extract_field_info(template_path)

    # 转换为字段列表格式
    fields_data = save_fields_to_json(fields_info)
    fields_list = fields_data['字段列表']

    logger.info(f"已提取 {len(fields_list)} 个唯一字段\n")

    # ========== 第二步：生成回填数据 ==========
    fill_data = {}

    # 保存关键数据
    year_value = None
    month_value = None
    day_value = None
    monthly_income = None
    proof_date = None

    # 第一遍：生成年份、月份、日期
    logger.info("=== 第二阶段：生成证明日期 ===")
    for field in fields_list:
        field_name = field['字段名']
        field_type = detect_field_type(field_name)

        if field_type == 'year':
            year_value = generate_value_by_type(field_type, monthly_income, year_value, month_value)
            fill_data[field_name] = {
                'value': year_value,
                'full_match': field['完整标记'],
                'has_underline': field['有下划线'],
                'type': field_type,
                'location_type': field['位置类型'],
                'context': field['上下文']
            }
            logger.info(f"  {field_name}: {year_value}")

        elif field_type == 'month':
            month_value = generate_value_by_type(field_type, monthly_income, year_value, month_value)
            fill_data[field_name] = {
                'value': month_value,
                'full_match': field['完整标记'],
                'has_underline': field['有下划线'],
                'type': field_type,
                'location_type': field['位置类型'],
                'context': field['上下文']
            }
            logger.info(f"  {field_name}: {month_value}")

        elif field_type == 'day':
            day_value = generate_value_by_type(field_type, monthly_income, year_value, month_value)
            fill_data[field_name] = {
                'value': day_value,
                'full_match': field['完整标记'],
                'has_underline': field['有下划线'],
                'type': field_type,
                'location_type': field['位置类型'],
                'context': field['上下文']
            }
            logger.info(f"  {field_name}: {day_value}")

    # 组合证明日期
    if year_value and month_value and day_value:
        proof_date = datetime(int(year_value), int(month_value), int(day_value))
        logger.info(f"\n证明日期: {proof_date.strftime('%Y-%m-%d')}\n")

    # 第二遍：生成其他字段
    logger.info("=== 第三阶段：生成其他字段 ===")

    # 提前生成月收入值
    monthly_income = fake.random_int(min=5000, max=50000)
    logger.info(f"  [预先生成] 月收入基数: {monthly_income}元")

    for field in fields_list:
        field_name = field['字段名']
        field_type = detect_field_type(field_name)

        if field_name in fill_data:
            continue

        if field_type == 'date':
            if proof_date:
                days_ago = fake.random_int(min=365, max=1825)
                start_date = proof_date - timedelta(days=days_ago)
                field_value = start_date.strftime('%Y-%m-%d')
                logger.info(f"  {field_name}: {field_value} (证明日期前{days_ago//30}个月)")
            else:
                field_value = generate_value_by_type(field_type, monthly_income, year_value, month_value)
                logger.info(f"  {field_name}: {field_value}")
        elif field_type == 'income_lowercase':
            field_value = f"{monthly_income}元"
            logger.info(f"  {field_name}: {field_value}")
        elif field_type == 'income_uppercase':
            field_value = number_to_chinese(monthly_income)
            logger.info(f"  {field_name}: {field_value}")
        else:
            field_value = generate_value_by_type(field_type, monthly_income, year_value, month_value)
            if field_type not in ['annual_income_uppercase', 'annual_income_lowercase']:
                logger.info(f"  {field_name}: {field_value}")

        fill_data[field_name] = {
            'value': field_value,
            'full_match': field['完整标记'],
            'has_underline': field['有下划线'],
            'type': field_type,
            'location_type': field['位置类型'],
            'context': field['上下文']
        }

    # 第三阶段：计算工作月数并更新年收入
    logger.info("\n=== 第四阶段：计算年收入 ===")
    hire_date_value = fill_data.get('入职日期', {}).get('value')
    if hire_date_value and proof_date and monthly_income:
        work_months = calculate_work_months(hire_date_value, proof_date.strftime('%Y-%m-%d'))
        logger.info(f"入职日期: {hire_date_value}")
        logger.info(f"证明日期: {proof_date.strftime('%Y-%m-%d')}")
        logger.info(f"工作月数: {work_months}个月")
        logger.info(f"月收入: {monthly_income}元")

        bonus = fake.random_int(min=0, max=50000)
        logger.info(f"年终奖: {bonus}元")

        annual_income = monthly_income * work_months + bonus
        logger.info(f"年收入: {annual_income}元 (= {monthly_income} × {work_months} + {bonus})")

        annual_income_lower = f"{annual_income}元"
        annual_income_upper = number_to_chinese(annual_income)

        if '年收入（小写）' in fill_data:
            fill_data['年收入（小写）']['value'] = annual_income_lower
            logger.info(f"\n[更新] 年收入（小写）: {annual_income_lower}")

        if '年收入（大写）' in fill_data:
            fill_data['年收入（大写）']['value'] = annual_income_upper
            logger.info(f"[更新] 年收入（大写）: {annual_income_upper}")

    logger.info(f"\n生成回填数据完成，共 {len(fill_data)} 个字段")

    return fill_data


# ============================================================
# 第三部分：文档填充功能
# ============================================================

def get_nested_table(table, table_path):
    """
    根据table_path获取嵌套表格
    table_path: [0] 表示第0个表格内的嵌套表格
    """
    current_table = table
    for idx in table_path:
        # 遍历当前表格的所有单元格，查找嵌套表格
        nested_found = False
        for row in current_table.rows:
            for cell in row.cells:
                if cell.tables:
                    # 找到嵌套表格
                    current_table = cell.tables[idx]
                    nested_found = True
                    break
            if nested_found:
                break
        if not nested_found:
            return None
    return current_table


def replace_in_paragraph(paragraph, full_match, value, has_underline=False):
    """
    在段落中替换文本（支持跨多个run的替换）

    Args:
        paragraph: 段落对象
        full_match: 要替换的完整标记，如 {{姓名}}
        value: 替换的值
        has_underline: 是否有下划线（需要特殊处理）
    """
    # 获取段落完整文本
    para_text = paragraph.text

    # 检查是否包含要替换的文本
    if full_match not in para_text:
        return False

    # 替换文本
    if has_underline:
        # 有下划线的字段特殊处理
        # 回填时保留花括号：{{value}}
        # 计算长度差值：回填内容长度 - 标记字段长度
        length_diff = len(value) - len(full_match)

        if length_diff > 0:
            # 需要删除多余的空格
            # 判断value主要包含中文还是数字/字母
            import re
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', value))
            total_chars = len(value)

            # 如果中文字符占比超过50%，则认为是中文内容，需要乘以2
            if total_chars > 0 and chinese_chars / total_chars > 0.5:
                multiplier = 2
            else:
                multiplier = 1

            # 查找标记字段后面的空格
            match_index = para_text.find(full_match)
            after_match = para_text[match_index + len(full_match):]

            # 计算需要删除的空格数
            spaces_to_remove = length_diff * multiplier + 4

            # 删除后面的空格，保留花括号
            new_text = para_text[:match_index] + "{{" + value + "}}" + after_match.lstrip()
            # 如果原文本后面有空格，需要删除一部分
            if len(after_match) > len(after_match.lstrip()):
                original_spaces = len(after_match) - len(after_match.lstrip())
                remaining_spaces = max(0, original_spaces - spaces_to_remove)
                new_text = para_text[:match_index] + "{{" + value + "}}" + " " * remaining_spaces + after_match.lstrip()
        else:
            # 差值不为正，正常替换，保留花括号
            new_text = para_text.replace(full_match, "{{" + value + "}}")
    else:
        # 没有下划线，直接替换
        new_text = para_text.replace(full_match, value)

    # 重新构建runs：将所有文本合并到第一个run中
    if paragraph.runs:
        # 清空第一个run的文本并设置为替换后的完整文本
        paragraph.runs[0].text = new_text

        # 清空其他runs的文本
        for run in paragraph.runs[1:]:
            run.text = ''

    return True


def replace_in_body(doc, paragraph_index, full_match, value, has_underline=False):
    """
    在文档主体段落中替换
    """
    if paragraph_index < len(doc.paragraphs):
        paragraph = doc.paragraphs[paragraph_index]
        return replace_in_paragraph(paragraph, full_match, value, has_underline)
    return False


def replace_in_nested_table(doc, context, full_match, value, has_underline=False):
    """
    在嵌套表格中替换（在所有单元格中搜索匹配的段落）
    table_path: [0] 表示表格[0]内的嵌套表格
    """
    table_path = context.get('table_path', [])

    # 使用get_nested_table函数获取嵌套表格
    for table in doc.tables:
        nested_table = get_nested_table(table, table_path)
        if nested_table:
            # 在嵌套表格的所有单元格中搜索包含full_match的段落
            for cell in nested_table._cells:
                for paragraph in cell.paragraphs:
                    if full_match in paragraph.text:
                        return replace_in_paragraph(paragraph, full_match, value, has_underline)
    return False


def add_underline_to_brackets(doc):
    """
    为文档中所有 {{value}} 格式的字段及其左右空格添加下划线
    """
    import re

    def process_paragraph(paragraph):
        """处理段落中的 {{value}} 格式"""
        para_text = paragraph.text
        # 查找所有 {{value}} 模式
        pattern = re.compile(r'\{\{[^}]+\}\}')
        matches = list(pattern.finditer(para_text))

        if not matches:
            return

        # 构建需要下划线的内容位置集合
        underline_ranges = []
        for match in matches:
            match_start = match.start()
            match_end = match.end()

            # 向前查找空格
            temp_pos = match_start - 1
            while temp_pos >= 0 and para_text[temp_pos] == ' ':
                temp_pos -= 1
            underline_start = temp_pos + 1

            # 向后查找空格
            temp_pos = match_end
            while temp_pos < len(para_text) and para_text[temp_pos] == ' ':
                temp_pos += 1
            underline_end = temp_pos

            underline_ranges.append((underline_start, underline_end))

        # 如果没有需要下划线的范围，直接返回
        if not underline_ranges:
            return

        # 记录原有 runs 的字体信息
        # 构建: 位置 -> 字体名称
        position_fonts = {}
        current_pos = 0
        for run in paragraph.runs:
            run_text = run.text
            run_len = len(run_text)

            # 记录这个 run 覆盖的所有位置的字体
            for i in range(run_len):
                pos = current_pos + i
                position_fonts[pos] = run.font.name

            current_pos += run_len

        # 获取某个位置的字体（处理跨 run 的情况）
        def get_font_for_position(pos):
            """获取指定位置的字体"""
            if pos in position_fonts:
                return position_fonts[pos]
            # 如果位置没有记录，尝试找最近的
            keys = sorted(position_fonts.keys())
            for key in keys:
                if key >= pos:
                    return position_fonts[key]
            # 如果都找不到，返回第一个位置的字体
            if keys:
                return position_fonts[keys[0]]
            # 默认字体
            return None

        # 清空所有原有 runs
        for run in paragraph.runs:
            run.text = ''

        # 按照下划线范围重新构建 runs
        current_pos = 0
        runs_to_add = []

        # 先将所有下划线范围按位置排序并合并
        underline_ranges.sort()
        merged_ranges = []
        for start, end in underline_ranges:
            if not merged_ranges:
                merged_ranges.append([start, end])
            else:
                _, last_end = merged_ranges[-1]
                if start <= last_end:
                    # 重叠或相邻，合并
                    merged_ranges[-1][1] = max(last_end, end)
                else:
                    merged_ranges.append([start, end])

        # 根据 merged_ranges 构建新的 runs，同时记录每个文本段对应的字体
        for start, end in merged_ranges:
            # 添加下划线范围前的文本（无下划线）
            if current_pos < start:
                before_text = para_text[current_pos:start]
                if before_text:
                    # 获取这个文本段的起始位置字体
                    font_name = get_font_for_position(current_pos)
                    runs_to_add.append((before_text, False, font_name))

            # 添加下划线范围的文本（有下划线）
            underline_text = para_text[start:end]
            if underline_text:
                # 获取这个文本段的起始位置字体
                font_name = get_font_for_position(start)
                runs_to_add.append((underline_text, True, font_name))
            current_pos = end

        # 添加剩余的文本（无下划线）
        if current_pos < len(para_text):
            remaining_text = para_text[current_pos:]
            if remaining_text:
                # 获取这个文本段的起始位置字体
                font_name = get_font_for_position(current_pos)
                runs_to_add.append((remaining_text, False, font_name))

        # 重新创建 runs，应用记录的字体
        if paragraph.runs and runs_to_add:
            # 使用第一个 run
            text, has_underline, font_name = runs_to_add[0]
            paragraph.runs[0].text = text.replace('{', ' ').replace('}', ' ')
            paragraph.runs[0].underline = has_underline
            if font_name:
                paragraph.runs[0].font.name = font_name

            # 添加其他 runs
            for text, has_underline, font_name in runs_to_add[1:]:
                run = paragraph.add_run(text.replace('{', ' ').replace('}', ' '))
                run.underline = has_underline
                if font_name:
                    run.font.name = font_name

    # 处理正文段落
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # 处理表格（包括嵌套表格）
    def process_tables(tables):
        for table in tables:
            for cell in table._cells:
                # 处理单元格段落
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
                # 处理嵌套表格
                if cell.tables:
                    process_tables(cell.tables)

    process_tables(doc.tables)

    print("\n已为所有 {{value}} 格式的字段添加下划线")


def fill_document(template_path, fill_data, output_path):
    """
    填充Word文档
    :param template_path: 模板文档路径
    :param fill_data: 填充数据字典
    :param output_path: 输出文档路径
    """
    # 使用提供的填充数据
    print(f"\n使用填充数据进行文档填充")

    # 打开模板文档
    print(f"正在打开模板文档: {template_path}")
    doc = Document(template_path)

    success_count = 0
    failed_count = 0
    failed_fields = []

    # 遍历所有字段进行替换
    for field_name, field_info in fill_data.items():
        value = field_info.get('value', '')
        full_match = field_info.get('full_match', '')
        location_type = field_info.get('location_type', 'body')
        context = field_info.get('context', {})
        has_underline = field_info.get('has_underline', False)

        print(f"\n处理字段: {field_name}")
        print(f"  替换内容: {full_match} -> {value}")
        print(f"  位置类型: {location_type}")
        print(f"  有下划线: {has_underline}")

        success = False

        if location_type == 'body':
            # 在文档主体中替换
            paragraph_index = context.get('paragraph_index', 0)
            success = replace_in_body(doc, paragraph_index, full_match, value, has_underline)
            if success:
                print(f"  [OK] 成功替换 (段落 {paragraph_index})")
            else:
                print(f"  [FAIL] 替换失败")

        elif location_type == 'nested_table':
            # 在嵌套表格中替换
            success = replace_in_nested_table(doc, context, full_match, value, has_underline)
            if success:
                print(f"  [OK] 成功替换 (嵌套表格)")

        if success:
            success_count += 1
        else:
            failed_count += 1
            failed_fields.append(field_name)

    # 为所有 {{value}} 格式的字段添加下划线
    print("\n正在为字段添加下划线...")
    add_underline_to_brackets(doc)

    # 保存文档
    print(f"\n正在保存文档到: {output_path}")

    # 如果输出文件已存在，先删除
    if os.path.exists(output_path):
        try:
            os.remove(output_path)
            print(f"已删除旧文件: {output_path}")
        except PermissionError:
            print(f"错误: 文件 {output_path} 正被其他程序占用，请先关闭该文件")
            raise

    doc.save(output_path)

    # 输出统计信息
    print("\n" + "="*50)
    print("填充完成！")
    print(f"成功替换: {success_count} 个字段")
    print(f"替换失败: {failed_count} 个字段")

    if failed_fields:
        print(f"\n失败的字段列表:")
        for field in failed_fields:
            print(f"  - {field}")

    print("="*50)

    return success_count, failed_count


# ============================================================
# 第四部分：主函数
# ============================================================

def process_template(template_path=None, output_doc_path=None):
    """
    处理模板文件：提取字段 → 生成回填数据 → 填充文档

    Args:
        template_path: 模板文档路径（可选，默认使用 template/income-template.docx）
        output_doc_path: 输出Word文档路径（可选，默认使用 output/01_words/income-template_filled.docx）

    Returns:
        tuple: (success_count, failed_count, company_name, fill_data) 成功和失败的字段数量、公司名称、本次填充的数据字典
    """
    # 获取脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # 如果未指定模板路径，使用默认路径
    if template_path is None:
        template_path = os.path.join(script_dir, 'template', 'income-template.docx')

    # 如果未指定输出文档路径，使用默认路径
    if output_doc_path is None:
        output_doc_path = os.path.join(script_dir, 'output', '01_words', 'income-template_filled.docx')

    print("="*60)
    print("收入证明文档处理工具")
    print("="*60)
    print(f"模板文档: {template_path}")
    print(f"输出文档: {output_doc_path}")

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_doc_path), exist_ok=True)

    # 第一步：提取字段并生成回填数据
    print("\n>>> 第一阶段：提取字段并生成回填数据")
    fill_data = generate_fill_data(template_path)

    # 第二步：使用回填数据填充文档
    print("\n>>> 第二阶段：填充文档")
    success_count, failed_count = fill_document(template_path, fill_data, output_doc_path)

    # 提取公司名称
    company_name = None
    if '公司名称' in fill_data:
        company_name = fill_data['公司名称']['value']
    elif '_company_name' in fill_data:
        company_name = fill_data['_company_name']['value']

    return success_count, failed_count, company_name, fill_data


def main():
    """主函数"""
    import argparse
    import sys

    parser = argparse.ArgumentParser(
        description='收入证明文档处理工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  # 使用默认路径
  python 1_process1.py

  # 指定模板路径
  python 1_process1.py template.docx

  # 指定输出文档路径
  python 1_process1.py --output-doc output.docx

  # 指定所有路径
  python 1_process1.py template.docx --output-doc output.docx

默认路径说明:
  - 默认模板: template/income-template.docx
  - 默认输出文档: output/01_words/income-template_filled.docx
        """
    )
    parser.add_argument('template_path', nargs='?', help='模板文档路径（可选，默认使用 template/income-template.docx）')
    parser.add_argument('--output-doc', help='输出Word文档路径（可选，默认使用 output/01_words/income-template_filled.docx）')

    args = parser.parse_args()

    try:
        success_count, failed_count, company_name, fill_data = process_template(
            template_path=args.template_path,
            output_doc_path=args.output_doc
        )

        print(f"\n最终统计: 成功 {success_count} 个，失败 {failed_count} 个")

        if company_name:
            print(f"公司名称: {company_name}")

        # 打印本次填充的数据
        print("\n" + "="*60)
        print("本次填充的数据:")
        print("="*60)
        print(json.dumps(fill_data, ensure_ascii=False, indent=2))
        print("="*60)

        if failed_count == 0:
            print("\n[SUCCESS] 所有字段处理成功!")
        else:
            print(f"\n[WARNING] 有 {failed_count} 个字段处理失败，请检查")

    except FileNotFoundError as e:
        print(f"错误: 找不到文件 - {e}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"错误: JSON文件解析失败 - {e}")
        sys.exit(1)
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
