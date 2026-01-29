"""
Word文档转图片工具
支持多种转换方式，自动选择最佳方案
"""

import os
import subprocess
import tempfile
import uuid
from pathlib import Path
from typing import Optional

from loguru import logger


class WordToImageConverter:
    """Word文档转图片转换器 - 支持多种转换策略"""

    def __init__(self, output_dir: str = "generated_images", dpi: int = 200):
        """
        初始化转换器

        Args:
            output_dir: 输出目录
            dpi: 图片DPI（推荐150-300）
        """
        self.output_dir = output_dir
        self.dpi = dpi
        Path(output_dir).mkdir(parents=True, exist_ok=True)

        # 检测可用的转换方法
        self.available_methods = self._detect_methods()
        self.best_method = self._select_best_method()

    def _detect_methods(self) -> dict:
        """
        检测所有可用的转换方法

        Returns:
            方法字典: {'method_name': bool}
        """
        methods = {
            'win32com_poppler': False,  # Word + PDF + poppler (最佳)
            'win32com_direct': False,   # Word直接渲染
            'docx2pdf_poppler': False,  # docx2pdf + poppler
            'python_docx': False,       # 纯Python方式
        }

        # 优先检测 LibreOffice (soffice)
        try:
            from shutil import which
            if which('soffice') or which('libreoffice'):
                methods['libreoffice_poppler'] = True
        except Exception:
            pass

        # 检测 win32com
        try:
            import win32com
            methods['win32com_poppler'] = True
            methods['win32com_direct'] = True
        except ImportError:
            pass

        # 检测 docx2pdf
        try:
            import docx2pdf
            methods['docx2pdf_poppler'] = True
        except ImportError:
            pass

        # 检测 python-docx (纯Python方式)
        try:
            import docx
            from PIL import Image, ImageDraw, ImageFont
            methods['python_docx'] = True
        except ImportError:
            pass

        # 检测 poppler
        self.poppler_path = self._find_poppler()

        return methods

    def _find_poppler(self) -> Optional[str]:
        """
        查找 poppler 安装路径

        Returns:
            poppler路径或None
        """
        # 常见安装路径
        possible_paths = [
            # 用户自定义（优先级最高）
            r"D:\software\Release-24.08.0-0\poppler-24.08.0\Library\bin",
            # 旧版本（避免使用）
            r"D:\software\Release-25.12.0-0\poppler-25.12.0\Library\bin",
            # 项目路径
            r"D:\DataProject\DataSynthesis\poppler-24.08.0\Library\bin",
            r"D:\DataProject\DataSynthesis\poppler-25.12.0\Library\bin",
            # Windows 常见位置
            r"C:\Program Files\poppler-24.08.0\Library\bin",
            r"C:\Program Files\poppler\Library\bin",
            r"C:\poppler-24.08.0\Library\bin",
            r"C:\poppler\Library\bin",
            # Linux/macOS
            "/usr/bin",
            "/usr/local/bin",
        ]

        for path in possible_paths:
            if os.path.exists(path):
                # 检查是否有 pdftoppm
                pdftoppm = os.path.join(path, "pdftoppm.exe" if os.name == 'nt' else "pdftoppm")
                if os.path.exists(pdftoppm):
                    return path

        return None

    def _select_best_method(self) -> str:
        """
        选择最佳转换方法

        优先级:
        1. win32com_poppler - 最佳质量
        2. docx2pdf_poppler - 不需要Word
        3. win32com_direct - Word直接渲染
        4. python_docx - 纯Python（备用）
        """
        # 优先使用 LibreOffice + poppler（跨平台，推荐）
        if self.available_methods.get('libreoffice_poppler') and self.poppler_path:
            return 'libreoffice_poppler'
        if self.available_methods.get('win32com_poppler') and self.poppler_path:
            return 'win32com_poppler'
        elif self.available_methods.get('docx2pdf_poppler') and self.poppler_path:
            return 'docx2pdf_poppler'
        elif self.available_methods.get('win32com_direct'):
            return 'win32com_direct'
        elif self.available_methods.get('python_docx'):
            return 'python_docx'
        else:
            return 'none'

    def convert(self, docx_path: str, output_name: Optional[str] = None) -> Optional[str]:
        """
        转换Word文档为图片

        Args:
            docx_path: Word文档路径
            output_name: 输出文件名（不含扩展名）

        Returns:
            图片路径，失败返回None
        """
        if not os.path.exists(docx_path):
            logger.info(f"错误: 文件不存在 {docx_path}")
            return None

        if output_name is None:
            output_name = Path(docx_path).stem

        output_path = os.path.join(self.output_dir, f"{output_name}.png")

        logger.info(f"\n转换方法: {self._get_method_name(self.best_method)}")

        try:
            if self.best_method == 'libreoffice_poppler':
                return self._convert_libreoffice_poppler(docx_path, output_path)
            elif self.best_method == 'win32com_poppler':
                return self._convert_win32com_poppler(docx_path, output_path)
            elif self.best_method == 'docx2pdf_poppler':
                return self._convert_docx2pdf_poppler(docx_path, output_path)
            elif self.best_method == 'win32com_direct':
                return self._convert_win32com_direct(docx_path, output_path)
            elif self.best_method == 'python_docx':
                return self._convert_python_docx(docx_path, output_path)
            else:
                logger.info("错误: 没有可用的转换方法")
                self._print_install_guide()
                return None
        except Exception as e:
            logger.info(f"转换失败: {e}")
            import traceback
            traceback.print_exc()
            return None

    def _get_method_name(self, method: str) -> str:
        """获取方法的中文名称"""
        names = {
            'win32com_poppler': 'Word → PDF → Poppler (最佳质量)',
            'docx2pdf_poppler': 'docx2pdf → Poppler',
            'win32com_direct': 'Word直接渲染',
            'python_docx': '纯Python渲染',
        }
        return names.get(method, method)

    def _convert_win32com_poppler(self, docx_path: str, output_path: str) -> str:
        """方法1: Word → PDF → Poppler (推荐，最佳质量)"""
        import win32com.client
        import time

        docx_path_abs = os.path.abspath(docx_path)
        output_path_abs = os.path.abspath(output_path)
        pdf_path = output_path_abs.replace('.png', '.pdf')

        word = None
        doc = None

        try:
            # 启动Word
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            # 打开文档
            doc = word.Documents.Open(docx_path_abs)
            time.sleep(0.5)

            # 保存为PDF
            doc.SaveAs(pdf_path, FileFormat=17)
            time.sleep(0.5)

            doc.Close(SaveChanges=0)

            # PDF转图片
            result = self._pdf_to_image(pdf_path, output_path_abs, docx_path_abs)

            # 清理PDF
            if os.path.exists(pdf_path):
                os.remove(pdf_path)

            return result

        except Exception as e:
            logger.info(f"  Word转PDF失败: {e}")
            # 降级到纯Python方法
            if os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                except:
                    pass
            return self._convert_python_docx(docx_path, output_path)
        finally:
            if doc:
                try:
                    doc.Close(SaveChanges=0)
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass

    def _convert_docx2pdf_poppler(self, docx_path: str, output_path: str) -> str:
        """方法2: docx2pdf → Poppler"""
        from docx2pdf import convert

        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = os.path.join(temp_dir, "temp.pdf")
            convert(docx_path, pdf_path)
            return self._pdf_to_image(pdf_path, output_path, docx_path)

    def _convert_libreoffice_poppler(self, docx_path: str, output_path: str) -> str:
        """使用 LibreOffice 将 docx 转为 PDF，然后用 pdf2image + poppler 转图片（推荐方案）"""
        # 尝试通过系统上的 soffice 命令转换
        from shutil import which

        soffice = which('soffice') or which('libreoffice')
        if not soffice:
            # 在 Windows 上可能存在安装路径，需要用户提供或使用 win32com 退回
            raise RuntimeError('未找到 LibreOffice 的 soffice 可执行文件')

        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                subprocess.run([
                    soffice,
                    f"-env:UserInstallation=file:///tmp/LibreOffice_Conversion_{str(uuid.uuid4())}",
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', temp_dir,
                    docx_path
                ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            except subprocess.CalledProcessError as e:
                # 转换失败，抛出异常以触发降级
                raise RuntimeError(f'LibreOffice 转 PDF 失败: {e.stderr.decode(errors="ignore")}')

            # 找到生成的 PDF
            pdfs = list(Path(temp_dir).glob('*.pdf'))
            if not pdfs:
                raise RuntimeError('LibreOffice 未生成 PDF 文件')
            pdf_path = str(pdfs[0])

            # 使用已有的 _pdf_to_image（会使用 self.poppler_path）
            return self._pdf_to_image(pdf_path, output_path, docx_path)

    def _convert_python_docx(self, docx_path: str, output_path: str) -> str:
        """方法4: 纯Python方式（备用）"""
        from docx import Document
        from PIL import Image, ImageDraw, ImageFont

        # 加载文档
        doc = Document(docx_path)

        # 计算图片尺寸
        max_width = 0
        total_height = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    total_height += run.font.size * 1.5  # 行高
                else:
                    total_height += 12  # 默认行高
                # 估算宽度
                max_width = max(max_width, len(run.text) * 10)

        # 创建图片
        image = Image.new("RGB", (max_width, int(total_height)), "white")
        draw = ImageDraw.Draw(image)

        # 绘制文本
        y = 0
        for para in doc.paragraphs:
            for run in para.runs:
                font_size = run.font.size if run.font.size else 12
                font = ImageFont.truetype("arial.ttf", int(font_size))
                draw.text((0, y), run.text, fill="black", font=font)
                y += int(font_size * 1.5)  # 行高

        # 保存图片
        image.save(output_path)

        return output_path

    def _pdf_to_image(self, pdf_path: str, image_path: str, docx_path: str) -> str:
        """
        PDF 转 图片

        Args:
            pdf_path: PDF文件路径
            image_path: 输出图片路径
            docx_path: 原始DOCX文件路径（用于错误报告）

        Returns:
            图片路径
        """
        from pdf2image import convert_from_path

        # 使用 poppler_path
        poppler_path = self.poppler_path

        # 转换 PDF 为图片
        try:
            images = convert_from_path(pdf_path, dpi=self.dpi, poppler_path=poppler_path)
            if not images:
                raise RuntimeError("PDF 转换未生成任何图片")

            # 保存第一张图片
            first_page = images[0]
            first_page.save(image_path, "PNG")

            return image_path
        except Exception as e:
            raise RuntimeError(f"PDF 转图片失败: {e}") from e

    def _print_install_guide(self):
        """打印安装指南"""
        logger.info(
            """
            安装指南:

            1. 请确保已安装以下软件:
               - Microsoft Word
               - LibreOffice
               - Poppler

            2. 确保将它们的安装路径添加到系统环境变量中.

            3. 推荐使用管理员权限运行此程序, 以避免权限问题.

            4. 如遇到问题, 请查看各软件的官方文档或社区寻求帮助.
            """
        )
