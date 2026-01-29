# -*- coding: utf-8 -*-

"""
Description:
    贷款报告模板填充器 - 将JSON数据填充到Word模板生成报告
Create: 2025/01/28
"""

import json
import os
import re
from typing import Dict, Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from loguru import logger

from datamate.core.base_op import Mapper


class LoanReportFiller(Mapper):
    """贷款报告模板填充器"""

    def __init__(self, *args, **kwargs):
        super(LoanReportFiller, self).__init__(*args, **kwargs)

        # 确保输出目录存在
        self._output_dir = None
        self._template_path = None

    def _load_json_data(self, json_file_path: str) -> Dict[str, Any]:
        """加载JSON数据文件"""
        with open(json_file_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def _fill_template(self, template_path: str, data: Dict[str, Any]) -> Document:
        """填充模板数据"""
        doc = Document(template_path)
        field_values = {}

        # 扁平化所有数据到字典中
        # 基本信息
        field_values["经办机构"] = data["基本信息"]["经办机构"]
        field_values["填表日期"] = data["基本信息"]["填表日期"]
        field_values["经办客户经理"] = data["基本信息"]["经办客户经理"]
        field_values["电话"] = data["基本信息"]["电话"]

        # 贷款申请信息
        field_values["借款人姓名"] = data["贷款申请信息"]["借款人姓名"]
        field_values["借款人身份证号码"] = data["贷款申请信息"]["借款人身份证号码"]
        field_values["借款人配偶姓名"] = data["贷款申请信息"]["借款人配偶姓名"]
        field_values["借款人配偶身份证号码"] = data["贷款申请信息"]["借款人配偶身份证号码"]
        field_values["单笔金额(元)"] = data["贷款申请信息"]["单笔金额(元)"]
        field_values["单笔期限(月)"] = data["贷款申请信息"]["单笔期限(月)"]
        field_values["还款方式"] = data["贷款申请信息"]["还款方式"]
        field_values["期供计算期(月)"] = data["贷款申请信息"]["期供计算期(月)"]
        field_values["贷款种类"] = data["贷款申请信息"]["贷款种类"]
        field_values["申请贷款利率基准利率浮动"] = data["贷款申请信息"]["申请贷款利率基准利率浮动"]

        # 借款人信息
        field_values["职业性质"] = data["借款人信息"]["职业性质"]
        field_values["年龄(20-65岁)"] = data["借款人信息"]["年龄(20-65岁)"]
        field_values["婚姻状况"] = data["借款人信息"]["婚姻状况"]
        field_values["家庭现有住房套数"] = data["借款人信息"]["家庭现有住房套数"]
        field_values["是否本地常住"] = data["借款人信息"]["是否本地常住"]
        field_values["借款人及其配偶有无不良信用记录"] = data["借款人信息"]["借款人及其配偶有无不良信用记录"]

        # 偿还能力评估
        field_values["借款人工作单位/岗位"] = data["偿还能力评估"]["借款人工作单位/岗位"]
        field_values["借款人收入证明材料"] = data["偿还能力评估"]["借款人收入证明材料"]
        field_values["借款人月收入(元)"] = data["偿还能力评估"]["借款人月收入(元)"]
        field_values["收入认定过程"] = data["偿还能力评估"]["收入认定过程"]
        field_values["配偶工作单位/岗位"] = data["偿还能力评估"]["配偶工作单位/岗位"]
        field_values["配偶收入证明材料"] = data["偿还能力评估"]["配偶收入证明材料"]
        field_values["配偶月收入(元)"] = data["偿还能力评估"]["配偶月收入(元)"]
        field_values["统计期间"] = data["偿还能力评估"]["统计期间"]
        field_values["流入资金总笔数"] = data["偿还能力评估"]["流入资金总笔数"]
        field_values["流入总资金"] = data["偿还能力评估"]["流入总资金"]
        field_values["月平均流入资金"] = data["偿还能力评估"]["月平均流入资金"]
        field_values["本笔贷款月供支出(元)"] = data["偿还能力评估"]["本笔贷款月供支出(元)"]
        field_values["家庭月收入总额(元)"] = data["偿还能力评估"]["家庭月收入总额(元)"]
        field_values["本笔贷款所购房产物业费(元)"] = data["偿还能力评估"]["本笔贷款所购房产物业费(元)"]
        field_values["本笔收入负债比(DTI≤50%)"] = data["偿还能力评估"]["本笔收入负债比(DTI≤50%)"]
        field_values["家庭其余贷款月负债(不含本笔,元)"] = data["偿还能力评估"]["家庭其余贷款月负债(不含本笔,元)"]
        field_values["总收入负债比(DTI≤55%)"] = data["偿还能力评估"]["总收入负债比(DTI≤55%)"]
        field_values["家庭月债总额"] = data["偿还能力评估"]["家庭月债总额"]

        # 风险抵押详细信息
        field_values["LTV(申请金额/押品评估价值)"] = data["风险抵押详细信息"]["LTV(申请金额/押品评估价值)"]
        field_values[" LTV(申请金额/押品评估价值)"] = data["风险抵押详细信息"]["LTV(申请金额/押品评估价值)"]  # 带空格的版本
        field_values["抵押物现使用状态"] = data["风险抵押详细信息"]["抵押物现使用状态"]
        field_values["房产名称及座别(座落地址)"] = data["风险抵押详细信息"]["房产名称及座别(座落地址)"]
        field_values["建筑面积(平方米)"] = data["风险抵押详细信息"]["建筑面积(平方米)"]
        field_values["楼龄(年)"] = data["风险抵押详细信息"]["楼龄(年)"]
        field_values["评估价值(元)"] = data["风险抵押详细信息"]["评估价值(元)"]
        field_values["交易价格(元)"] = data["风险抵押详细信息"]["交易价格(元)"]
        field_values["评估单价(元/平方米)"] = data["风险抵押详细信息"]["评估单价(元/平方米)"]
        field_values["认定价值(元)"] = data["风险抵押详细信息"]["认定价值(元)"]

        # 调查结论
        field_values["借款人基础资料真实齐全，贷款申请事实清晰、真实、有效"] = data["调查结论"]["借款人基础资料真实齐全，贷款申请事实清晰、真实、有效"]
        field_values["贷款用途合法合规，贷款担保物权属关系清晰"] = data["调查结论"]["贷款用途合法合规，贷款担保物权属关系清晰"]
        field_values["抵押物估值符合规定，抵押率符合审批要求"] = data["调查结论"]["抵押物估值符合规定，抵押率符合审批要求"]
        field_values["借款人及其配偶征信良好，无重大风险"] = data["调查结论"]["借款人及其配偶征信良好，无重大风险"]
        field_values["经调查本住房贷款为借款人家庭自住使用"] = data["调查结论"]["经调查本住房贷款为借款人家庭自住使用"]
        field_values["整体结论"] = data["调查结论"]["整体结论"]
        field_values["其它"] = ""

        # 表格填充（查找占位符格式：{{字段名}}）
        if doc.tables:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        for field_name, field_value in field_values.items():
                            placeholder = f"{{{{{field_name}}}}}"
                            if placeholder in cell_text:
                                new_text = cell_text.replace(placeholder, str(field_value))
                                cell.text = new_text
                                # 设置字体
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = '微软雅黑'
                                        run.font.size = Pt(8)
                                        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 段落填充（查找占位符格式：{{字段名}}）
        for i, paragraph in enumerate(doc.paragraphs):
            current_text = paragraph.text.strip()
            for field_name, field_value in field_values.items():
                placeholder = f"{{{{{field_name}}}}}"
                if placeholder in current_text:
                    new_text = current_text.replace(placeholder, str(field_value))
                    paragraph.clear()
                    run = paragraph.add_run(new_text)
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(8)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        return doc

    def _parse_filename(self, filename: str) -> tuple:
        """从文件名提取sequence_number和borrower_name"""
        # 文件名格式：虚拟贷款调查报告_0000001-杨勇.json
        match = re.match(r'虚拟贷款调查报告_(\d{7})-(.+)\.json', filename)
        if match:
            return match.group(1), match.group(2)
        return None, None

    def _process_single_file(self, json_file: str) -> bool:
        """处理单个JSON文件"""
        filename = os.path.basename(json_file)
        logger.info(f"处理文件: {filename}")

        # 加载数据
        data = self._load_json_data(json_file)
        if data is None:
            return False

        # 填充模板
        doc = self._fill_template(self._template_path, data)
        if doc is None:
            return False

        # 解析文件名
        sequence_number, borrower_name = self._parse_filename(filename)
        if sequence_number is None:
            borrower_name = data["贷款申请信息"]["借款人姓名"]
            sequence_number = "0000000"

        # 生成输出文件名
        output_filename = f"个人贷款调查报告_{sequence_number}-{borrower_name}.docx"
        output_path = os.path.join(self._output_dir, output_filename)

        # 保存文档
        doc.save(output_path)
        logger.info(f"报告已生成: {output_path}")

        return True

    def execute(self, sample: Dict[str, Any]) -> Dict[str, Any]:
        """执行模板填充"""

        file_path = sample.get('filePath')
        if not file_path.endswith('.docx') or os.path.normpath(file_path).count(os.sep) > 3:
            return sample

        self._template_path = sample["filePath"]
        self._output_dir = sample["export_path"]

        # 检查模板文件
        if not os.path.exists(self._template_path):
            logger.error(f"模板文件不存在: {self._template_path}")
            return sample

        # 获取输入数据目录
        input_dir = self._output_dir
        if not os.path.exists(input_dir):
            logger.error(f"数据目录不存在: {input_dir}")
            return sample

        # 获取所有JSON文件
        json_files = [f for f in os.listdir(input_dir) if f.endswith('.json')]
        json_files = [os.path.join(input_dir, f) for f in json_files]

        if not json_files:
            logger.warning(f"在目录 {input_dir} 中未找到JSON文件")
            return sample

        logger.info(f"找到 {len(json_files)} 个JSON文件待处理")

        # 批量处理
        success_count = 0
        for json_file in json_files:
            if self._process_single_file(json_file):
                success_count += 1

        logger.info(f"处理完成: 成功 {success_count}/{len(json_files)}")

        return sample
