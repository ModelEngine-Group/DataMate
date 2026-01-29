"""
图片转换模块
负责将docx文档转换为图片
"""

import os
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional
import shutil


class ImageConverter:
    """图片转换器 - 将docx转换为图片"""

    def __init__(self, output_dir: str = "generated_images", dpi: int = 200, instance_id = ""):
        """
        初始化图片转换器

        Args:
            output_dir: 输出目录
            dpi: 图片DPI（影响清晰度）
        """
        self.output_dir = output_dir
        self.dpi = dpi
        Path(output_dir).mkdir(parents=True, exist_ok=True)

        # 检测可用的转换方法
        self.method = self._detect_conversion_method()
        self.instance_id = instance_id

    def _detect_conversion_method(self) -> str:
        """
        检测可用的转换方法

        Returns:
            'win32com' | 'libreoffice' | 'docx2pdf' | 'none'
        """
        # 检查是否在Windows上且有win32com
        try:
            import win32com
            return "win32com"
        except ImportError:
            pass

        # 检查是否有LibreOffice
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/usr/bin/libreoffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]
        for path in libreoffice_paths:
            if os.path.exists(path):
                return "libreoffice"

        # 检查docx2pdf
        try:
            import docx2pdf
            return "docx2pdf"
        except ImportError:
            pass

        return "none"

    def convert_single(self, docx_path: str, output_name: Optional[str] = None) -> Optional[str]:
        """
        转换单个docx文件为图片

        Args:
            docx_path: docx文件路径
            output_name: 输出文件名（不含扩展名），默认使用docx文件名

        Returns:
            生成的图片路径，失败返回None
        """
        if output_name is None:
            output_name = Path(docx_path).stem

        output_path = os.path.join(self.output_dir, f"{output_name}.png")

        try:
            if self.method == "win32com":
                return self._convert_with_win32com(docx_path, output_path)
            elif self.method == "libreoffice":
                return self._convert_with_libreoffice(docx_path, output_path)
            elif self.method == "docx2pdf":
                return self._convert_with_docx2pdf(docx_path, output_path)
            else:
                print(f"错误: 没有可用的转换方法，请安装win32com或LibreOffice")
                return None
        except Exception as e:
            print(f"转换失败 {docx_path}: {e}")
            return None

    def _convert_with_win32com(self, docx_path: str, output_path: str) -> str:
        """使用win32com转换（Windows专用）- Word → PDF → 图片"""
        import win32com.client
        import time

        # 转换为绝对路径（win32com需要）
        docx_path_abs = os.path.abspath(docx_path)
        output_path_abs = os.path.abspath(output_path)
        pdf_path = output_path_abs.replace('.png', '.pdf')

        word = None
        doc = None

        try:
            # 启动Word应用（使用DispatchEx创建独立实例，避免批量转换时冲突）
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            # 打开文档
            doc = word.Documents.Open(docx_path_abs)
            time.sleep(0.5)  # 等待文档完全加载

            # 保存为PDF（FileFormat=17 是PDF格式）
            doc.SaveAs(pdf_path, FileFormat=17)
            time.sleep(0.5)  # 等待保存完成

            doc.Close(SaveChanges=0)
            time.sleep(0.2)

            # 将PDF转换为图片（使用Poppler）
            img_path = self._pdf_to_image(pdf_path, output_path_abs, docx_path_abs)

            # 清理临时PDF文件
            if os.path.exists(pdf_path):
                os.remove(pdf_path)

            return output_path_abs

        except Exception as e:
            # 失败时使用备用方法
            print(f"  PDF转换失败: {e}，使用备用方法...")
            return self._fallback_render_from_docx(docx_path_abs, output_path_abs)
        finally:
            if doc:
                try:
                    doc.Close(SaveChanges=0)
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass

    def _word_to_image_direct(self, docx_path: str, output_path: str) -> str:
        """
        直接使用Word将文档渲染为图片
        通过截图Word窗口内容来实现
        """
        import win32com.client
        import win32gui
        from PIL import ImageGrab
        import time

        word = None
        doc = None

        try:
            # 启动Word
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = True
            word.DisplayAlerts = 0
            word.WindowState = 1  # 最大化

            time.sleep(0.5)  # 等待Word启动

            # 打开文档
            doc = word.Documents.Open(os.path.abspath(docx_path))

            # 设置视图为打印视图
            doc.ActiveWindow.View.Type = 1  # wdPrintView
            doc.ActiveWindow.View.Zoom.Percentage = 100

            # 滚动到顶部确保内容可见
            doc.ActiveWindow.VerticalPercentScrolled = 0

            # 更长的等待时间，确保文档完全渲染
            time.sleep(2.0)

            # 获取窗口句柄 - 尝试多种方式
            try:
                hwnd = int(word.ActiveWindow.Hwnd)
            except:
                # 如果无法获取 Hwnd，通过窗口标题查找
                def callback(hwnd, windows):
                    class_name = win32gui.GetClassName(hwnd)
                    if class_name == 'OpusApp':
                        windows.append(hwnd)
                    return True

                windows = []
                win32gui.EnumWindows(callback, windows)
                if windows:
                    hwnd = windows[0]
                else:
                    raise RuntimeError("无法找到Word窗口")

            if not hwnd:
                raise RuntimeError("无法获取Word窗口句柄")

            # 将窗口置于前台
            win32gui.SetForegroundWindow(hwnd)
            time.sleep(0.3)

            # 获取窗口矩形
            left, top, right, bottom = win32gui.GetWindowRect(hwnd)

            print(f"  窗口位置: ({left}, {top}) -> ({right}, {bottom})")

            # 获取客户区（文档内容区域）
            # 尝试直接获取客户区位置
            try:
                # 左上角坐标
                client_left, client_top = win32gui.ClientToScreen(hwnd, (0, 0))
                # 客户区大小
                client_rect = win32gui.GetClientRect(hwnd)
                client_right = client_left + client_rect[2]
                client_bottom = client_top + client_rect[3]

                print(f"  客户区: ({client_left}, {client_top}) -> ({client_right}, {client_bottom})")

                # 截图客户区
                screenshot = ImageGrab.grab(bbox=(client_left, client_top, client_right, client_bottom))

            except Exception as e:
                print(f"  客户区截图失败: {e}，使用全窗口截图")
                # 备选：截图整个窗口
                screenshot = ImageGrab.grab(bbox=(left, top, right, bottom))

            # 保存为PNG
            screenshot.save(output_path, 'PNG')

            print(f"  图片已保存: {output_path}")

            # 关闭文档
            doc.Close(SaveChanges=0)

            return output_path

        finally:
            if doc:
                try:
                    doc.Close(SaveChanges=0)
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass

    def _convert_with_libreoffice(self, docx_path: str, output_path: str) -> str:
        """使用LibreOffice转换"""
        import tempfile

        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            # 转换为PDF
            subprocess.run([
                "libreoffice",
                f"-env:UserInstallation=file:///tmp/LibreOffice_Conversion_{self.instance_id}",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", temp_dir,
                docx_path
            ], check=True, capture_output=True)

            # 找到生成的PDF文件
            pdf_file = Path(temp_dir) / f"{Path(docx_path).stem}.pdf"

            if pdf_file.exists():
                return self._pdf_to_image(str(pdf_file), output_path)
            else:
                raise Exception("LibreOffice转换失败，未生成PDF文件")

    def _convert_with_docx2pdf(self, docx_path: str, output_path: str) -> str:
        """使用docx2pdf转换"""
        import tempfile
        from docx2pdf import convert

        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = os.path.join(temp_dir, "temp.pdf")
            convert(docx_path, pdf_path)
            return self._pdf_to_image(pdf_path, output_path)

    def _pdf_to_image(self, pdf_path: str, output_path: str, docx_path: str = None) -> str:
        """将PDF转换为图片"""
        try:
            from pdf2image import convert_from_path
            from PIL import Image

            # 转换PDF为图片列表，指定 poppler 路径
            images = convert_from_path(pdf_path, dpi=self.dpi)

            if len(images) == 1:
                # 单页，直接保存
                images[0].save(output_path, 'PNG')
            else:
                # 多页，保存第一页或合并
                images[0].save(output_path, 'PNG')

            return output_path

        except Exception as e:
            # 如果pdf2image失败，使用备用方法
            print(f"\n  PDF转图片失败: {e}")
            if docx_path:
                return self._fallback_render_from_docx(docx_path, output_path)
            else:
                return self._fallback_render_from_docx(pdf_path.replace('.pdf', '.docx'), output_path)

    def _fallback_render_from_docx(self, docx_path: str, output_path: str) -> str:
        """
        备用方法：直接从docx渲染简单图片（不需要poppler）
        注意：此方法不保留完美格式，仅用于测试
        """
        from docx import Document
        from PIL import Image, ImageDraw, ImageFont

        # 读取docx内容
        doc = Document(docx_path)

        # 创建白色背景图片
        width = 800
        line_height = 25
        padding = 40
        height = padding * 2

        # 计算需要的总高度
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    all_text.append(row_text)

        height += len(all_text) * line_height

        # 创建图片
        img = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)

        # 尝试使用系统字体
        try:
            font = ImageFont.truetype("msyh.ttc", 16)  # 微软雅黑
            font_bold = ImageFont.truetype("msyhbd.ttc", 18)
        except:
            try:
                font = ImageFont.truetype("arial.ttf", 16)
                font_bold = ImageFont.truetype("arialbd.ttf", 18)
            except:
                font = ImageFont.load_default()
                font_bold = font

        # 绘制文本
        y = padding
        for i, text in enumerate(all_text):
            # 根据内容选择字体
            if "贷款结清证明" in text or text.endswith("证明"):
                current_font = font_bold
            else:
                current_font = font

            draw.text((padding, y), text, fill='black', font=current_font)
            y += line_height

        # 保存图片
        img.save(output_path, 'PNG')
        return output_path

    def convert_batch(self, docx_paths: List[str]) -> List[str]:
        """
        批量转换docx文件

        Args:
            docx_paths: docx文件路径列表

        Returns:
            生成的图片路径列表
        """
        output_paths = []

        print(f"\n使用方法: {self.method}")
        print(f"开始转换 {len(docx_paths)} 个文档...")

        for i, docx_path in enumerate(docx_paths):
            print(f"  [{i+1}/{len(docx_paths)}] 转换: {Path(docx_path).name}", end="")

            output_name = Path(docx_path).stem
            try:
                result = self.convert_single(docx_path, output_name)
                if result:
                    output_paths.append(result)
                    print(" -> 成功")
                else:
                    print(" -> 失败")
            except Exception as e:
                print(f" -> 失败: {e}")

        print(f"\n成功转换 {len(output_paths)}/{len(docx_paths)} 个文件")
        return output_paths

    def convert_directory(self, input_dir: str, pattern: str = "*.docx") -> List[str]:
        """
        转换整个目录的docx文件

        Args:
            input_dir: 输入目录
            pattern: 文件匹配模式

        Returns:
            生成的图片路径列表
        """
        input_path = Path(input_dir)
        docx_files = list(input_path.glob(pattern))

        print(f"在 {input_dir} 中找到 {len(docx_files)} 个docx文件")

        return self.convert_batch([str(f) for f in docx_files])


class SimpleImageConverter:
    """
    简化版图片转换器
    使用docx2pdf + pdf2image的方案
    适用于没有win32com的环境
    """

    def __init__(self, output_dir: str = "generated_images"):
        self.output_dir = output_dir
        Path(output_dir).mkdir(parents=True, exist_ok=True)

    def convert(self, docx_path: str, output_name: Optional[str] = None) -> Optional[str]:
        """
        简单的docx到png转换
        需要安装: pip install docx2pdf pdf2image
        Windows还需要安装poppler
        """
        if output_name is None:
            output_name = Path(docx_path).stem

        output_path = os.path.join(self.output_dir, f"{output_name}.png")

        try:
            # 方法1: docx2pdf + pdf2image
            from docx2pdf import convert
            from pdf2image import convert_from_path
            import tempfile

            with tempfile.TemporaryDirectory() as temp_dir:
                pdf_path = os.path.join(temp_dir, "temp.pdf")

                # docx转pdf
                convert(docx_path, pdf_path)

                # pdf转图片
                images = convert_from_path(pdf_path, dpi=200)

                if images:
                    images[0].save(output_path, 'PNG')
                    return output_path

        except Exception as e:
            print(f"转换失败: {e}")
            print("\n请确保安装了以下依赖:")
            print("  pip install docx2pdf pdf2image")
            print("Windows还需要下载poppler并添加到PATH")
            print("  下载地址: https://github.com/oschwartz10612/poppler-windows/releases/")

        return None


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        # 命令行使用: python image_converter.py <docx_file_or_directory>
        converter = ImageConverter(output_dir="generated_images")

        input_path = sys.argv[1]

        if os.path.isfile(input_path):
            result = converter.convert_single(input_path)
            if result:
                print(f"\n已转换: {result}")
        elif os.path.isdir(input_path):
            results = converter.convert_directory(input_path)
            print(f"\n共转换 {len(results)} 个文件")
    else:
        # 测试代码
        print("=" * 60)
        print("图片转换器测试")
        print("=" * 60)

        converter = ImageConverter()

        print(f"\n检测到的转换方法: {converter.method}")

        if converter.method == "none":
            print("\n没有可用的转换方法！")
            print("请安装以下依赖之一:")
            print("  1. pip install pywin32 (Windows)")
            print("  2. 安装LibreOffice")
            print("  3. pip install docx2pdf pdf2image")
        else:
            print("\n转换器已就绪，可以转换docx文件")
