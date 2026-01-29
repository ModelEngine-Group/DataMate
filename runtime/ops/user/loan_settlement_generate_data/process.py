import os
import json
import re
from typing import Dict, Any, List, Optional
from pathlib import Path
from docx import Document
from datamate.core.base_op import Mapper
from .src import DataGenerator
from loguru import logger

# ---------------------------------------------------------------------
# 导入说明：
# 请确保您的 DataGenerator 类可以被导入。
# 如果您将 src 文件夹也打包进来了，请使用: from .src.data_generator import DataGenerator
# 这里假设您将 data_generator.py 放到了同级目录:
# from .data_generator import DataGenerator
# 为防止报错，此处仅做占位，请根据实际打包结构调整 import
# ---------------------------------------------------------------------

class LoanSettlementDataGenOperator(Mapper):
    """
    数据生成算子：DataGenOperator
    """
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # 获取 metadata.yml 中的参数
        self.count = int(kwargs.get('countParam', 5))
        
        # 处理 seed，前端 input 传过来可能是字符串
        seed_val = kwargs.get('seedParam', 42)
        self.seed = int(seed_val) if seed_val else None
        
        self.output_dir = kwargs.get('outputDirParam', '/dataset').strip()

    def _extract_fields_from_text(self, text: str) -> List[str]:
        """从文本中提取 {{字段名}} 格式的字段"""
        pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(pattern, text)
        return [match.strip() for match in matches]

    def _extract_placeholder_fields(self, template_path: str) -> List[str]:
        """从 Word 模板中提取字段"""
        doc = Document(template_path)
        fields_set = set()

        # 在段落中查找
        for para in doc.paragraphs:
            fields_set.update(self._extract_fields_from_text(para.text))

        # 在表格中查找
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        fields_set.update(self._extract_fields_from_text(para.text))

        return list(fields_set)
    
    def generate_complete_records(self, field_names: List[str], count: int, seed: Optional[int] = None) -> List[Dict[str, Any]]:
        """
        生成完整的数据记录（正确处理字段间的依赖关系）

        Args:
            field_names: 字段名列表
            count: 生成数量
            seed: 随机种子

        Returns:
            生成的完整数据记录列表
        """
        generator = DataGenerator(seed=seed)
        records = []

        for _ in range(count):
            # 使用 DataGenerator 的 generate_single_record 方法
            # 这个方法已经处理了所有依赖关系
            full_record = generator.generate_single_record()

            # 只保留模板中需要的字段
            record = {}
            for field_name in field_names:
                if field_name in full_record:
                    record[field_name] = full_record[field_name]

            records.append(record)

        return records

    def execute(self, sample: Dict[str, Any]) -> Dict[str, Any]:
        """
        核心执行逻辑
        :param sample: 输入样本（生成算子通常忽略输入，或将其作为触发上下文）
        """
        file_path = sample.get('filePath')
        if not file_path.endswith('.docx') or os.path.normpath(file_path).count(os.sep) > 3:
            return sample

        try:
            # 这里的 import 放在 execute 内部或文件头均可，取决于您的打包方式
            # from .data_generator import DataGenerator 
            # 为了演示，假设 DataGenerator 可以在此处被调用
            
            # 1. 解析模板
            field_names = self._extract_placeholder_fields(sample['filePath'])
            
            # 2. 生成数据 (调用外部依赖逻辑)
            # 注意：此处需要您确保 DataGenerator 可用
            records = self.generate_complete_records(field_names, self.count, self.seed)
            
            # -----------------------------------------------------------

            # 3. 如果配置了输出目录，则保存文件
            if self.output_dir:
                output_path = Path(str(sample.get('export_path')))
                output_path.mkdir(parents=True, exist_ok=True)
                
                # 将文件路径写入 sample 返回
                sample['text'] = json.dumps(records, ensure_ascii=False, indent=2)
            
        except Exception as e:
            logger.error(f"Error in DataGenOperator: {e}")
            
        return sample