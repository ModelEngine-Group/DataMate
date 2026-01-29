"""
Word模板填充算子
功能：将不动产数据填充到Word模板中，生成查询结果文档
"""

import os
import copy
import json
import base64
import tempfile
from typing import Dict, Any
from loguru import logger

from datamate.core.base_op import Mapper


class WordFiller:
    """Word模板填充辅助类"""

    def __init__(self, template_path: str, output_dir: str):
        """
        初始化Word填充器

        Args:
            template_path: Word模板文件路径
            output_dir: 填充后文档输出目录
        """
        self.template_path = template_path
        self.output_dir = output_dir

        # 确保输出目录存在
        os.makedirs(self.output_dir, exist_ok=True)

        logger.info(f"Word填充器初始化完成，模板路径: {template_path}")

    def _replace_text_simple(
        self, element, replacements: Dict[str, str], bold: bool = False
    ):
        """
        简单的文本替换方法（直接替换整个元素的文本）
        注意：如果段落不包含占位符，则保留原始格式（包括边框）

        Args:
            element: 段落或单元格对象
            replacements: 替换字典 {占位符: 替换值}
            bold: 是否加粗
        """
        # 检查是否有段落属性
        if not hasattr(element, "runs"):
            # 如果是单元格，遍历其段落
            if hasattr(element, "paragraphs"):
                for para in element.paragraphs:
                    self._replace_text_simple(para, replacements, bold)
            return

        # 获取当前文本
        current_text = element.text

        # 检查是否需要替换（如果没有任何占位符，则跳过）
        needs_replacement = False
        for key in replacements.keys():
            if key in current_text:
                needs_replacement = True
                break

        # 如果不需要替换，保留原始格式，直接返回
        if not needs_replacement:
            return

        # 替换文本
        new_text = current_text
        for key, value in replacements.items():
            new_text = new_text.replace(key, str(value))

        # 清空元素
        element.clear()

        # 添加新文本并设置加粗
        run = element.add_run(new_text)

        # 设置加粗属性
        run.bold = bold

    def _find_table_with_placeholder(self, doc, placeholder: str) -> int:
        """
        查找包含指定占位符的表格

        Args:
            doc: Word文档对象
            placeholder: 占位符

        Returns:
            表格索引，未找到返回-1
        """
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        return table_idx
        return -1

    def _find_placeholder_row(self, table, placeholder: str) -> int:
        """
        查找包含占位符的行

        Args:
            table: 表格对象
            placeholder: 占位符

        Returns:
            行索引，未找到返回-1
        """
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if placeholder in cell.text:
                    return row_idx
        return -1

    def _clone_table_row(self, table, source_row_idx: int) -> int:
        """
        克隆表格中的一行

        Args:
            table: 表格对象
            source_row_idx: 要克隆的行索引

        Returns:
            新行的索引
        """
        source_row = table.rows[source_row_idx]

        # 获取源行的XML元素
        source_tr = source_row._tr

        # 深度克隆XML元素
        new_tr = copy.deepcopy(source_tr)

        # 插入新行到表格末尾
        table._tbl.append(new_tr)

        return len(table.rows) - 1

    def _set_paragraph_bold(self, paragraph, is_bold: bool = True):
        """
        设置段落中所有文本的加粗状态

        Args:
            paragraph: 段落对象
            is_bold: 是否加粗
        """
        for run in paragraph.runs:
            run.bold = is_bold

    def _set_cell_bold(self, cell, is_bold: bool = True):
        """
        设置单元格中所有文本的加粗状态

        Args:
            cell: 单元格对象
            is_bold: 是否加粗
        """
        for paragraph in cell.paragraphs:
            self._set_paragraph_bold(paragraph, is_bold)

    def _get_province_without_suffix(self, province: str) -> str:
        """
        获取用于标题的省份名称（直辖市和普通省份都保留完整名称）

        Args:
            province: 省份名称（如"河北省"、"北京市"）

        Returns:
            用于标题的省份名称（直辖市返回完整名称如"北京市"，普通省份返回完整名称如"河北省"）
        """
        # 直接返回完整名称（保留"省"字）
        return province

    def fill_single_document(self, data: Dict[str, Any], output_filename: str) -> str:
        """
        填充单个Word文档

        Args:
            data: 填充数据字典
            output_filename: 输出文件名

        Returns:
            输出文件路径
        """
        try:
            from docx import Document
            from docx.table import Table

            # 每次都从原始模板重新加载
            doc = Document(self.template_path)

            # 获取去掉后缀的省份名称（用于标题）
            province = data.get("省份", "")
            province_without_suffix = self._get_province_without_suffix(province)

            # 准备替换字典 - 段落级别的占位符
            paragraph_replacements = {
                "{{省份}}": province_without_suffix,  # 使用去掉后缀的省份名称
                "{{姓名（证件号码）}}": data.get("姓名（证件号码）", ""),
                "{{查询目的}}": data.get("查询目的", ""),
                "{{姓名}}": data.get("姓名（证件号码）", "").split("（")[0]
                if "（" in data.get("姓名（证件号码）", "")
                else data.get("姓名（证件号码）", ""),
                "{{电脑查册人}}": data.get("电脑查册人", ""),
                "{{证明流水号}}": data.get("证明流水号", ""),
                "{{打印日期}}": data.get("打印日期", ""),
            }

            # 替换段落中的占位符，并设置加粗
            for paragraph in doc.paragraphs:
                # 检查是否是标题段落（包含"不动产信息查询结果"）
                is_title = "不动产信息查询结果" in paragraph.text

                # 替换文本，如果是标题段落则加粗
                self._replace_text_simple(
                    paragraph, paragraph_replacements, bold=is_title
                )

            # 处理表格数据
            table_data = data.get("表格数据", [])

            # 准备表格级别的替换数据（包括姓名、查询目的等）
            table_replacements = {
                "{{姓名（证件号码）}}": data.get("姓名（证件号码）", ""),
                "{{查询目的}}": data.get("查询目的", ""),
                "{{姓名}}": data.get("姓名（证件号码）", "").split("（")[0]
                if "（" in data.get("姓名（证件号码）", "")
                else data.get("姓名（证件号码）", ""),
                "{{电脑查册人}}": data.get("电脑查册人", ""),
                "{{证明流水号}}": data.get("证明流水号", ""),
                "{{打印日期}}": data.get("打印日期", ""),
            }

            # 处理所有表格中的占位符
            for table_idx, table in enumerate(doc.tables):
                # 先替换表格中的所有占位符（包括姓名、查询目的等）
                for row_idx, row in enumerate(table.rows):
                    for cell in row.cells:
                        # 检查是否是表格1的第一行（列名行）
                        is_table1_header = table_idx == 0 and row_idx == 0
                        # 检查是否是表格3（所有单元格都加粗）
                        is_table3 = table_idx == 2

                        # 替换文本
                        self._replace_text_simple(cell, table_replacements)

                        # 设置加粗
                        if is_table1_header or is_table3:
                            self._set_cell_bold(cell, True)

            # 查找包含表格占位符的表格（用于处理动态表格行）
            table_idx = self._find_table_with_placeholder(doc, "{{序号}}")

            if table_idx >= 0 and table_data:
                table = doc.tables[table_idx]

                # 找到占位符行
                placeholder_row_idx = self._find_placeholder_row(table, "{{序号}}")

                if placeholder_row_idx >= 0:
                    # 先克隆所有需要的行（在填充之前）
                    # 第一行使用原有的占位符行
                    # 后续行克隆占位符行
                    for i in range(1, len(table_data)):
                        self._clone_table_row(table, placeholder_row_idx)

                    # 现在填充所有行数据
                    for i, row_data in enumerate(table_data):
                        # 目标行索引
                        target_row_idx = placeholder_row_idx + i

                        # 准备当前行的替换数据
                        row_replacements = {
                            "{{序号}}": str(row_data.get("序号", "")),
                            "{{状态}}": row_data.get("状态", ""),
                            "{{合同号}}": row_data.get("合同号", ""),
                            "{{地址}}": row_data.get("地址", ""),
                            "{{产权证号}}": row_data.get("产权证号", ""),
                            "{{规划用途}}": row_data.get("规划用途", ""),
                            "{{建筑面积}}": str(row_data.get("建筑面积", "")),
                            "{{数据来源}}": row_data.get("数据来源", ""),
                        }

                        # 替换目标行中的占位符
                        target_row = table.rows[target_row_idx]
                        for cell in target_row.cells:
                            self._replace_text_simple(cell, row_replacements)

            # 保存填充后的文档
            output_path = os.path.join(self.output_dir, output_filename)
            doc.save(output_path)

            logger.info(f"文档填充完成，保存到: {output_path}")
            return output_path

        except FileNotFoundError:
            logger.error(f"模板文件未找到: {self.template_path}")
            raise
        except Exception as e:
            logger.error(f"填充文档时发生错误: {str(e)}")
            raise


class WordTemplateFiller(Mapper):
    """
    Word模板填充算子
    类名建议使用驼峰命名法定义，例如 WordTemplateFiller
    """

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # 从UI参数获取配置
        self.filler = None
        self.template_path = kwargs.get(
            "templatePath", "不动产信息查询结果模板.docx"
        )
        self.data_source = kwargs.get("dataSource", "input")
        self.fixed_data = kwargs.get("fixedData", "")
        self.output_format = kwargs.get("outputFormat", "word")

        # 创建临时输出目录
        self.temp_output_dir = None



    def execute(self, sample: Dict[str, Any]) -> Dict[str, Any]:
        """
        核心处理逻辑
        :param sample: 输入的数据样本，通常包含 text_key 等字段
        :return: 处理后的数据样本
        """
        # 创建Word填充器实例
        self.template_path = sample["filePath"]
        self.filler = WordFiller(self.template_path, sample["export_path"])
        try:
            # 解析数据
            if self.data_source == "input":
                # 从输入文本解析JSON数据
                input_text = sample.get(self.text_key, "")
                try:
                    data_list = json.loads(input_text)
                    if not isinstance(data_list, list):
                        data_list = [data_list] if data_list else []
                except json.JSONDecodeError:
                    logger.error(f"无法解析输入文本为JSON: {input_text[:100]}")
                    raise ValueError("输入文本不是有效的JSON格式")
            else:
                # 使用固定数据
                try:
                    data_list = json.loads(self.fixed_data)
                    if not isinstance(data_list, list):
                        data_list = [data_list]
                except json.JSONDecodeError:
                    logger.error(f"无法解析固定数据为JSON: {self.fixed_data[:100]}")
                    raise ValueError("固定数据不是有效的JSON格式")

            if not data_list or len(data_list) == 0:
                logger.warning("没有数据需要填充")
                sample[self.text_key] = "没有数据需要填充"
                return sample

            # 循环填充文档，将文件保存到对应路径
            output_paths = []
            for idx, data in enumerate(data_list):
                output_filename = f"不动产查询表_{data.get('index', idx + 1):03d}.docx"
                output_path = self.filler.fill_single_document(data, output_filename)
                output_paths.append(output_path)

            logger.info(f"成功填充{len(output_paths)}个Word文档，格式: {self.output_format}")
            return sample

        except Exception as e:
            logger.error(f"填充Word模板时发生错误: {str(e)}")
            raise
