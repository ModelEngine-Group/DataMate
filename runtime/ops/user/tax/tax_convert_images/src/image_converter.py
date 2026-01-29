"""
图片转换模块 - 将docx文档转换为图片
"""

import os
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional


class ImageConverter:
    """图片转换器 - 将docx转换为图片"""

    def __init__(self, output_dir: str, dpi: int = 200):
        """
        初始化图片转换器

        Args:
            output_dir: 输出目录（存放jpg图片）
            dpi: 图片DPI（影响清晰度）
        """
        self.output_dir = output_dir
        self.dpi = dpi
        Path(output_dir).mkdir(parents=True, exist_ok=True)

        # 检测可用的转换方法
        self.method = self._detect_conversion_method()

    def _detect_conversion_method(self) -> str:
        """
        检测可用的转换方法

        Returns:
            'win32com' | 'libreoffice' | 'docx2pdf' | 'none'
        """
        # 检查是否在Windows上且有win32com
        try:
            import win32com
            return "win32com"
        except ImportError:
            pass

        # 检查是否有LibreOffice
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/usr/bin/libreoffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]
        for path in libreoffice_paths:
            if os.path.exists(path):
                return "libreoffice"

        # 检查docx2pdf
        try:
            import docx2pdf
            return "docx2pdf"
        except ImportError:
            pass

        return "none"

    def convert_single(self, docx_path: str, output_name: Optional[str] = None) -> Optional[str]:
        """
        转换单个docx文件为图片

        Args:
            docx_path: docx文件路径
            output_name: 输出文件名（不含扩展名），默认使用docx文件名

        Returns:
            生成的图片路径，失败返回None
        """
        if output_name is None:
            output_name = Path(docx_path).stem

        output_path = os.path.join(self.output_dir, f"{output_name}.jpg")

        try:
            if self.method == "win32com":
                return self._convert_with_win32com(docx_path, output_path)
            elif self.method == "libreoffice":
                return self._convert_with_libreoffice(docx_path, output_path)
            elif self.method == "docx2pdf":
                return self._convert_with_docx2pdf(docx_path, output_path)
            else:
                print(f"错误: 没有可用的转换方法，请安装win32com或LibreOffice")
                return None
        except Exception as e:
            print(f"转换失败 {docx_path}: {e}")
            return None

    def _convert_with_win32com(self, docx_path: str, output_path: str) -> str:
        """使用win32com转换（Windows专用）- Word → PDF → 图片"""
        import win32com.client
        import time

        # 转换为绝对路径（win32com需要）
        docx_path_abs = os.path.abspath(docx_path)
        output_path_abs = os.path.abspath(output_path)
        pdf_path = output_path_abs.replace('.jpg', '.pdf')

        word = None
        doc = None

        try:
            # 启动Word应用（使用DispatchEx创建独立实例，避免批量转换时冲突）
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            # 打开文档
            doc = word.Documents.Open(docx_path_abs)
            time.sleep(0.5)  # 等待文档完全加载

            # 保存为PDF（FileFormat=17 是PDF格式）
            doc.SaveAs(pdf_path, FileFormat=17)
            time.sleep(0.5)  # 等待保存完成

            doc.Close(SaveChanges=0)
            time.sleep(0.2)

            # 将PDF转换为图片（使用Poppler）
            img_path = self._pdf_to_image(pdf_path, output_path_abs, docx_path_abs)

            # 清理临时PDF文件
            if os.path.exists(pdf_path):
                os.remove(pdf_path)

            return output_path_abs

        except Exception as e:
            # 失败时使用备用方法
            print(f"  PDF转换失败: {e}，使用备用方法...")
            return self._fallback_render_from_docx(docx_path_abs, output_path_abs)
        finally:
            if doc:
                try:
                    doc.Close(SaveChanges=0)
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass

    def _convert_with_libreoffice(self, docx_path: str, output_path: str) -> str:
        """使用LibreOffice转换"""
        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            # 转换为PDF
            subprocess.run([
                "soffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", temp_dir,
                docx_path
            ], check=True, capture_output=True)

            # 找到生成的PDF文件
            pdf_file = Path(temp_dir) / f"{Path(docx_path).stem}.pdf"

            if pdf_file.exists():
                return self._pdf_to_image(str(pdf_file), output_path)
            else:
                raise Exception("LibreOffice转换失败，未生成PDF文件")

    def _convert_with_docx2pdf(self, docx_path: str, output_path: str) -> str:
        """使用docx2pdf转换"""
        from docx2pdf import convert

        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = os.path.join(temp_dir, "temp.pdf")
            convert(docx_path, pdf_path)
            return self._pdf_to_image(pdf_path, output_path)

    def _pdf_to_image(self, pdf_path: str, output_path: str, docx_path: str = None) -> str:
        """将PDF转换为图片"""
        try:
            from pdf2image import convert_from_path
            from PIL import Image

            # Poppler 路径 - 根据实际情况配置
            poppler_path = None  # 如果poppler在PATH中，不需要指定

            # 如果Windows系统，尝试使用常见路径
            if os.name == 'nt':
                possible_paths = [
                    r"D:\DataProject\DataSynthesis\poppler-25.12.0\Library\bin",
                    r"C:\Program Files\poppler\Library\bin",
                    r"C:\Program Files (x86)\poppler\Library\bin",
                ]
                for path in possible_paths:
                    if os.path.exists(path):
                        poppler_path = path
                        break

            # 转换PDF为图片列表
            if poppler_path:
                images = convert_from_path(pdf_path, dpi=self.dpi, poppler_path=poppler_path)
            else:
                images = convert_from_path(pdf_path, dpi=self.dpi)

            if len(images) == 1:
                # 单页，直接保存
                images[0].save(output_path, 'JPEG', quality=95)
            else:
                # 多页，保存第一页或合并
                images[0].save(output_path, 'JPEG', quality=95)

            return output_path

        except Exception as e:
            # 如果pdf2image失败，使用备用方法
            print(f"\n  PDF转图片失败: {e}")
            if docx_path:
                return self._fallback_render_from_docx(docx_path, output_path)
            else:
                return self._fallback_render_from_docx(pdf_path.replace('.pdf', '.docx'), output_path)

    def _fallback_render_from_docx(self, docx_path: str, output_path: str) -> str:
        """
        备用方法：直接从docx渲染简单图片（不需要poppler）
        注意：此方法不保留完美格式，仅用于测试
        """
        from docx import Document
        from PIL import Image, ImageDraw, ImageFont

        # 读取docx内容
        doc = Document(docx_path)

        # 创建白色背景图片
        width = 800
        line_height = 25
        padding = 40
        height = padding * 2

        # 计算需要的总高度
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    all_text.append(row_text)

        height += len(all_text) * line_height

        # 创建图片
        img = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)

        # 尝试使用系统字体
        try:
            font = ImageFont.truetype("msyh.ttc", 16)  # 微软雅黑
            font_bold = ImageFont.truetype("msyhbd.ttc", 18)
        except:
            try:
                font = ImageFont.truetype("arial.ttf", 16)
                font_bold = ImageFont.truetype("arialbd.ttf", 18)
            except:
                font = ImageFont.load_default()
                font_bold = font

        # 绘制文本
        y = padding
        for i, text in enumerate(all_text):
            # 根据内容选择字体
            if "完税证明" in text or text.endswith("证明"):
                current_font = font_bold
            else:
                current_font = font

            draw.text((padding, y), text, fill='black', font=current_font)
            y += line_height

        # 保存图片
        img.save(output_path, 'JPEG', quality=95)
        return output_path

    def convert_batch(self, docx_paths: List[str]) -> List[str]:
        """
        批量转换docx文件

        Args:
            docx_paths: docx文件路径列表

        Returns:
            生成的图片路径列表
        """
        output_paths = []

        print(f"\n使用方法: {self.method}")
        print(f"开始转换 {len(docx_paths)} 个文档...")

        for i, docx_path in enumerate(docx_paths):
            print(f"  [{i+1}/{len(docx_paths)}] 转换: {Path(docx_path).name}", end="")

            output_name = Path(docx_path).stem
            try:
                result = self.convert_single(docx_path, output_name)
                if result:
                    output_paths.append(result)
                    print(" -> 成功")
                else:
                    print(" -> 失败")
            except Exception as e:
                print(f" -> 失败: {e}")

        print(f"\n成功转换 {len(output_paths)}/{len(docx_paths)} 个文件")
        return output_paths

    def convert_directory(self, input_dir: str, pattern: str = "*.docx") -> List[str]:
        """
        转换整个目录的docx文件

        Args:
            input_dir: 输入目录
            pattern: 文件匹配模式

        Returns:
            生成的图片路径列表
        """
        input_path = Path(input_dir)
        docx_files = list(input_path.glob(pattern))

        print(f"在 {input_dir} 中找到 {len(docx_files)} 个docx文件")

        return self.convert_batch([str(f) for f in docx_files])


if __name__ == "__main__":
    # 测试代码
    converter = ImageConverter(output_dir="output/02_images")

    print(f"\n检测到的转换方法: {converter.method}")

    if converter.method == "none":
        print("\n没有可用的转换方法！")
        print("请安装以下依赖之一:")
        print("  1. pip install pywin32 (Windows)")
        print("  2. 安装LibreOffice")
        print("  3. pip install docx2pdf pdf2image")
