"""
文档生成模块 - 将数据填充到Word模板
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from typing import Dict
import os


class DocGenerator:
    """文档生成器 - 将数据填充到Word模板"""

    def __init__(self, template_path: str):
        """
        初始化文档生成器

        Args:
            template_path: Word模板文件路径
        """
        self.template_path = template_path

    def fill_tax_certificate(self, tax_data: Dict, output_path: str):
        """
        填充个人所得税完税证明模板

        参数:
        tax_data: 纳税人数据字典
        output_path: 输出文件路径
        """

        # 加载模板文档
        doc = Document(self.template_path)

        # 设置全局字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(10.5)

        # 填充纳税人信息表（第一个表格）
        if len(doc.tables) > 0:
            taxpayer_table = doc.tables[0]

            # 确保表格有正确的行数和列数
            if len(taxpayer_table.rows) >= 2 and len(taxpayer_table.columns) >= 4:
                # 第一行：填发日期
                taxpayer_table.cell(0, 3).text = tax_data.get('填发日期', '')
                # 第二行：纳税人姓名和身份证照类型
                taxpayer_table.cell(1, 1).text = tax_data.get('纳税人姓名', '')
                taxpayer_table.cell(1, 4).text = tax_data.get('纳税人身份证照类型', '')

                # 第三行：身份号码和凭证编号
                taxpayer_table.cell(2, 1).text = tax_data.get('纳税人身份号码', '')
                taxpayer_table.cell(2, 3).text = tax_data.get('凭证编码', '')

        # 填充纳税项目表（第二个表格）
        if len(doc.tables) > 1:
            tax_items_table = doc.tables[1]

            # 获取纳税项目数据
            tax_items = tax_data.get('纳税项目', [])

            # 填充纳税项目数据（从第二行开始，第一行为标题）
            for i, item in enumerate(tax_items):
                if i + 1 < len(tax_items_table.rows):
                    tax_items_table.cell(i + 1, 0).text = item.get('item', '')
                    tax_items_table.cell(i + 1, 1).text = item.get('period', '')
                    tax_items_table.cell(i + 1, 2).text = item.get('amount', '')

        # 填充收入合计表（第三个表格）
        if len(doc.tables) > 2:
            summary_table = doc.tables[2]

            # 确保表格有足够的行数
            if len(summary_table.rows) >= 5:
                # 工资薪金所得小计
                if len(summary_table.rows[0].cells) >= 2:
                    summary_table.cell(0, 1).text = tax_data.get('工资、薪金所得小计', '')

                # 劳务报酬所得
                if len(summary_table.rows) > 1 and len(summary_table.rows[1].cells) >= 2:
                    summary_table.cell(1, 1).text = tax_data.get('劳务报酬所得', '')

                # 稿酬所得
                if len(summary_table.rows) > 2 and len(summary_table.rows[2].cells) >= 2:
                    summary_table.cell(2, 1).text = tax_data.get('稿酬所得', '')

                # 税款金额合计（小写）
                if len(summary_table.rows) > 3 and len(summary_table.rows[3].cells) >= 2:
                    summary_table.cell(3, 1).text = tax_data.get('税款金额合计（小写）', '')

                # 税款金额合计（大写）
                if len(summary_table.rows) > 4 and len(summary_table.rows[4].cells) >= 2:
                    summary_table.cell(4, 1).text = tax_data.get('税款金额合计（大写）', '')

        # 保存输出文件
        doc.save(output_path)
        print(f"✅ 完税证明已成功填充并保存到: {output_path}")

    def batch_fill(self, data_list, output_dir: str):
        """
        批量填充多个数据到Word模板

        Args:
            data_list: 数据字典列表
            output_dir: 输出目录
        """
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)

        for i, tax_data in enumerate(data_list):
            # 生成安全的文件名（避免特殊字符）
            taxpayer_name = tax_data.get('纳税人姓名', f'纳税人{i + 1}')
            safe_name = "".join(c if c not in r'\/:*?"<>|' else "_" for c in taxpayer_name)

            # 构造输出路径
            output_path = os.path.join(output_dir, f"{safe_name}_个人所得税完税证明.docx")

            # 填充并保存
            self.fill_tax_certificate(tax_data, output_path)


if __name__ == "__main__":
    # 测试代码
    from pathlib import Path

    # 查找模板
    template_path = "../../template/个人所得税完税证明.docx"
    if not os.path.exists(template_path):
        # 尝试其他路径
        template_path = Path(__file__).parent.parent.parent / "template" / "个人所得税完税证明.docx"

    if not os.path.exists(template_path):
        print(f"❌ 找不到模板文件")
        exit(1)

    # 创建生成器
    generator = DocGenerator(str(template_path))

    # 测试数据
    test_data = {
        "填发日期": "填发日期：2025年01月25日",
        "纳税人姓名": "张三",
        "纳税人身份证照类型": "居民身份证",
        "纳税人身份号码": "11010119******1234",
        "凭证编码": "(2025)市区个证100号",
        "纳税项目": [
            {"item": "工资、薪金所得小计", "period": "2024年01月", "amount": "15700.00"},
            {"item": "劳务报酬所得", "period": "2024年06月", "amount": "2400.00"},
            {"item": "稿酬所得", "period": "2024年09月", "amount": "800.00"}
        ],
        "工资、薪金所得小计": "15700.00",
        "劳务报酬所得": "2400.00",
        "稿酬所得": "800.00",
        "税款金额合计（小写）": "18900.00",
        "税款金额合计（大写）": "壹万捌仟玖佰元整"
    }

    # 测试生成
    output_path = "output/01_words/test_个人所得税完税证明.docx"
    os.makedirs("output/01_words", exist_ok=True)
    generator.fill_tax_certificate(test_data, output_path)
